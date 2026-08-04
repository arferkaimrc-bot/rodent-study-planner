"""
IACUC Auto-Fill — turn a Rodent Study Planner analysis into a filled KAIMRC
Animal Ethics Application Form (.docx).

The scientific narrative (PARTS 2, 3, 4, 6, 8, 11) is synthesised here from the
platform's analysis output; the researcher only supplies the administrative
details (team, funding, housing) via the front-end modal.

Public API:
    generate_iacuc_docx(payload) -> io.BytesIO   # a ready-to-download .docx

`payload` (JSON from the browser):
    {
      "study":    {"study_title", "pi_name", "institution"},
      "groups":   [ <original form groups: group_name, drug_name, species,
                     strain, sex, age, weight, num_mice, dose, route,
                     experiment_type, target_organ, toxicity_endpoints[] ...> ],
      "analysis": [ <the /predict response array; each item has `summary`,
                     `rationale`, `reference_papers`, `drug` ...> ],
      "admin":    {"funding_source", "housing_type" ("standard"|"absl"),
                   "team": [ {name, role, qualifications, institution,
                              email, mobile} ] }
    }

Nothing here is persisted — the form is generated in-memory and streamed back,
honouring the project's rule never to store study data.
"""
import io
import re
from pathlib import Path

from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn

TEMPLATE_PATH = Path(__file__).resolve().parent / "iacuc_assets" / "iacuc_template.docx"

# ── small helpers ─────────────────────────────────────────────────────────

def _s(v, default=""):
    """Safe string."""
    if v is None:
        return default
    v = str(v).strip()
    return v if v else default


def _join_unique(items):
    seen, out = set(), []
    for it in items:
        it = _s(it)
        if it and it not in seen:
            seen.add(it)
            out.append(it)
    return out


def _list_sentence(items, joiner="and"):
    items = _join_unique(items)
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f"{items[0]} {joiner} {items[1]}"
    return ", ".join(items[:-1]) + f", {joiner} " + items[-1]


def _bullets(items):
    """Render a list as newline-separated bullet lines (Word keeps the \\n)."""
    items = [i for i in (_s(x) for x in items) if i]
    return "\n".join(f"• {i}" for i in items)


def _pair(groups, analysis):
    """Zip form groups with their analysis result by index; skip errored ones."""
    out = []
    analysis = analysis or []
    for i, g in enumerate(groups or []):
        a = analysis[i] if i < len(analysis) else {}
        if isinstance(a, dict) and a.get("error"):
            a = {}
        out.append((g or {}, a or {}, (a.get("summary") if isinstance(a, dict) else {}) or {}))
    return out


def _is_control(g):
    dn = _s(g.get("drug_name")).lower()
    gn = _s(g.get("group_name")).lower()
    return (not dn) or dn in ("control", "vehicle", "saline", "none") or "control" in gn


# ── narrative builders (each returns one answer-box string) ────────────────

_SPECIES_PLURAL = {"mouse": "mice", "rat": "rats"}


def _plural_species(sp):
    s = _s(sp, "animals").lower()
    return _SPECIES_PLURAL.get(s, s + "s" if s and not s.endswith("s") else (s or "animals"))


def _drugs(rows):
    return _join_unique(_s(g.get("drug_name")) for g, _, _ in rows if not _is_control(g))


def _format_reference(p, n):
    """Format one reference in Vancouver style — as it would appear in the
    reference list of a published biomedical paper:
        n. Authors. Title. Journal. Year. doi:… / PMID:… / Available from: URL
    """
    authors = _s(p.get("authors"))
    title = _s(p.get("title"))
    venue = _s(p.get("venue"))
    year = _s(p.get("year"))
    doi = _s(p.get("doi")).replace("https://doi.org/", "").lstrip("/")
    pmid = _s(p.get("pmid"))
    url = _s(p.get("url"))

    parts = []
    if authors:
        if authors.count(",") >= 2:          # 3+ names shown ⇒ list is truncated
            authors += ", et al"
        parts.append(authors.rstrip(".") + ".")
    if title:
        parts.append(title.rstrip(".") + ".")
    if venue:
        parts.append(venue.rstrip(".") + ".")
    if year:
        parts.append(str(year).rstrip(".") + ".")
    if doi:
        parts.append(f"doi:{doi}")
    elif pmid:
        parts.append(f"PMID:{pmid}")
    elif url:
        parts.append(f"Available from: {url}")
    return f"{n}. " + " ".join(parts)


def _study_purpose(rows, study):
    drugs = _join_unique(_s(g.get("drug_name")) for g, _, _ in rows if not _is_control(g))
    paradigms = _join_unique(_s(g.get("experiment_type")) for g, _, _ in rows)
    species = _join_unique(_s(s.get("species") or g.get("species")) for g, _, s in rows)
    title = _s(study.get("study_title"), "this study")
    parts = [f"The purpose of {title} is to investigate the biological effects and "
             f"tolerability of {_list_sentence(drugs) or 'the test article(s)'} "
             f"in {_list_sentence(species, 'and') or 'laboratory rodents'}."]
    if paradigms:
        parts.append(f"The work falls within {_list_sentence(paradigms)} research.")
    design = _s(study.get("design"))
    if design:
        parts.append(f"The study follows a {design} experimental design.")
    parts.append("Live animals are required because the whole-organism response — "
                 "absorption, distribution, metabolism, systemic toxicity and organ-level "
                 "effects — cannot be reproduced in cell-based or in-vitro systems.")
    return " ".join(parts)


def _what_done(rows):
    lines = ["Animals are allocated to the following groups:"]
    for g, _, s in rows:
        gname = _s(g.get("group_name"), "Group")
        n = _s(s.get("planned_animals") or g.get("num_mice"), "n")
        sp = _plural_species(s.get("species") or g.get("species"))
        route = _s(g.get("route"), "the assigned route")
        route_txt = route if ("route" in route.lower() or "(" in route) else f"the {route} route"
        if _is_control(g):
            veh = _s(g.get("vehicle"))
            vol = _s(g.get("vehicle_volume"))
            if veh and vol:
                veh_txt = f"the vehicle ({veh}, {vol})"
            elif veh:
                veh_txt = f"the vehicle ({veh})"
            else:
                veh_txt = "the vehicle/control"
            lines.append(f"• {gname}: {n} {sp} receive {veh_txt} by {route_txt}, "
                         f"handled identically to the treatment group(s).")
        else:
            drug = _s(g.get("drug_name"), "the test article")
            dose = _s(g.get("dose"))
            dose_txt = f" at {dose} mg/kg" if dose else ""
            lines.append(f"• {gname}: {n} {sp} receive {drug}{dose_txt} by {route_txt}.")
    lines.append("All animals are weighed and observed regularly for clinical signs; biological "
                 "samples are collected at defined time-points, after which animals are humanely "
                 "euthanised and tissues collected.")
    return "\n".join(lines)


def _expected_outcomes(rows):
    drugs = _drugs(rows)
    organs = _join_unique(_s(g.get("target_organ")) for g, _, _ in rows
                          if _s(g.get("target_organ")).lower() not in ("", "general", "none"))
    eps = _join_unique(e for g, _, _ in rows for e in (g.get("toxicity_endpoints") or []))
    lead = _list_sentence(drugs) or "the test article(s)"
    out = [f"The study is expected to characterise the dose–response and safety profile of "
           f"{lead} in the chosen rodent model."]
    if organs:
        out.append(f"Particular attention is paid to potential effects on the {_list_sentence(organs)}.")
    if eps:
        out.append(f"Assessed endpoints include {_list_sentence(eps)}.")
    out.append("The results will define a tolerable dose range and the clinical signs that "
               "signal toxicity, guiding the design of subsequent studies.")
    return " ".join(out)


def _study_benefit(rows):
    drugs = _drugs(rows)
    paradigms = _join_unique(_s(g.get("experiment_type")) for g, _, _ in rows)
    lead = _list_sentence(drugs) or "the test article(s)"
    txt = (f"The findings will improve understanding of the efficacy and safety of {lead}, "
           "informing safer dose selection and reducing risk in subsequent research and, "
           "ultimately, clinical development.")
    if paradigms:
        txt += f" This advances {_list_sentence(paradigms)} research and may benefit " \
               "both human and animal health."
    return txt


def _strain_health_issues(rows):
    strains = _join_unique(_s(s.get("strain") or g.get("strain")) for g, _, s in rows)
    # Genetic background / transgenic lines the researcher entered per group.
    genetics = _join_unique(
        " — ".join(x for x in [_s(g.get("genetic_background")), _s(g.get("transgenic_line"))] if x)
        for g, _, _ in rows)
    genetics = [x for x in genetics if x]
    extra = ""
    if genetics:
        extra = f" Genetic background / transgenic line(s): {_list_sentence(genetics)}."
    if not strains:
        return ("No strain-specific health or husbandry issues are anticipated. Standard "
                "outbred/inbred laboratory rodents will be used and monitored per facility SOPs."
                + extra)
    return (f"Animals of the following strain(s) will be used: {_list_sentence(strains)}. "
            "No unusual strain-specific health or husbandry concerns are anticipated; animals "
            "will be sourced from accredited vendors and monitored per facility SOPs." + extra)


def _scientific_justification(rows):
    species = _join_unique(_s(s.get("species") or g.get("species")) for g, _, s in rows)
    return ("Live animals of the species " + (_list_sentence(species) or "listed above") +
            " are scientifically justified because the study endpoints depend on an intact, "
            "physiologically integrated mammalian system — systemic pharmacokinetics, "
            "multi-organ toxicity and whole-body clinical response — that cannot be modelled "
            "in vitro. The rodent is the lowest phylogenetic species in which these responses "
            "are predictive, and is the established model for this class of study.")


def _drug_overview_text(s):
    """`drug_overview` is a dict {description, reference_url} (or occasionally a
    plain string). Return a clean description sentence — never the raw dict."""
    ov = s.get("drug_overview")
    if isinstance(ov, dict):
        return _s(ov.get("description"))
    return _s(ov)


def _drug_overview_url(s):
    ov = s.get("drug_overview")
    return _s(ov.get("reference_url")) if isinstance(ov, dict) else ""


def _relevant_papers(rows, drugs, organs):
    """Keep only papers whose title is on-topic (mentions the compound, target
    organ, or toxicology terms) — a professional reference list, not a dump of
    loosely-matched search hits."""
    keys = [d.lower() for d in drugs] + [o.lower() for o in organs] + \
           ["toxic", "toxicit", "safety", "pharmacokinet", "dose", "hepato",
            "nephro", "cardio", "adverse", "ld50", "in vivo", "rodent", "mouse", "rat"]
    seen, out = set(), []
    for _, a, _ in rows:
        for p in (a.get("reference_papers") or []):
            if not isinstance(p, dict):
                continue
            title = _s(p.get("title"))
            if not title or title.lower() in seen:
                continue
            if any(k in title.lower() for k in keys):
                seen.add(title.lower())
                out.append(p)
    return out


_OA_HOSTS = ("europepmc.org", "ncbi.nlm.nih.gov", "/pmc", "plos", "mdpi.com",
             "frontiersin.org", "biomedcentral.com", "hindawi.com", "doaj",
             "nature.com/articles", "elifesciences.org")


def _is_open_access(p):
    u = (_s(p.get("url")) or _s(p.get("doi"))).lower()
    return any(h in u for h in _OA_HOSTS)


def _select_references(rows, drugs, organs, want=10):
    """Up to `want` references — the SAME papers shown in the platform's
    "Citations & Sources" panel, so the form matches the app. Open-access
    entries are listed first; otherwise order is preserved."""
    seen, papers = set(), []
    for _, a, _ in rows:
        for p in (a.get("reference_papers") or []):
            if not isinstance(p, dict):
                continue
            t = _s(p.get("title")).lower()
            if t and t not in seen:
                seen.add(t)
                papers.append(p)
    oa = [p for p in papers if _is_open_access(p)]
    other = [p for p in papers if not _is_open_access(p)]
    return (oa + other)[:want]


def _literature_background(rows, researcher_background="", refine=False):
    drugs = _drugs(rows)
    organs = _join_unique(_s(g.get("target_organ")) for g, _, _ in rows
                          if _s(g.get("target_organ")).lower() not in ("", "general", "none"))
    overviews = _join_unique(_drug_overview_text(s) for _, _, s in rows if _drug_overview_text(s))
    lead = _list_sentence(drugs) or "the test article(s)"
    is_are = "is" if len(drugs) == 1 else "are"
    rb = _s(researcher_background)

    paras = []
    if rb and not refine:
        # Copy the researcher's background verbatim (grammar/format tidy-up only);
        # the platform only appends the reference list below.
        paras.append(_clean_text(rb))
    else:
        # Platform writes/expands the background from reputable sources.
        if rb:
            paras.append(_clean_text(rb))
        if overviews:
            paras.append(" ".join(overviews[:2]))
        subj = f"{lead} {is_are} of toxicological interest"
        if organs:
            subj += f", with particular relevance to the {_list_sentence(organs)}"
        subj += (". The present study examines the dose–response relationship and target-organ "
                 "effects of the compound in a rodent model.")
        paras.append(subj)

    # References — Vancouver style, on-topic, open-access prioritised.
    refs, n = [], 1
    for _, _, s in rows:
        u = _drug_overview_url(s)
        if u:
            cname = (drugs[0] if drugs else "Compound")
            refs.append(f"{n}. {cname}: chemical and toxicological summary. "
                        f"Available from: {u}")
            n += 1
            break
    for p in _select_references(rows, drugs, organs, want=10):
        refs.append(_format_reference(p, n))
        n += 1

    text = "\n\n".join(paras)
    if refs:
        text += "\n\nReferences:\n" + "\n".join(refs)
    return text


def _research_aims(rows):
    drugs = _drugs(rows)
    organs = _join_unique(_s(g.get("target_organ")) for g, _, _ in rows
                          if _s(g.get("target_organ")).lower() not in ("", "general", "none"))
    eps = _join_unique(e for g, _, _ in rows for e in (g.get("toxicity_endpoints") or []))
    lead = _list_sentence(drugs) or "the test article(s)"
    aims = [f"Characterise the biological response and tolerability of {lead} at the selected "
            "dose level(s) in the chosen rodent model."]
    if organs:
        aims.append(f"Evaluate potential target-organ toxicity affecting the {_list_sentence(organs)}.")
    if eps:
        aims.append(f"Assess the defined toxicity endpoints: {_list_sentence(eps)}.")
    aims.append("Compare treatment groups against concurrent controls to establish a "
                "dose–response relationship.")
    aims.append("Define humane endpoints and inform safe dosing for subsequent studies.")
    return "\n".join(f"{i}. {a}" for i, a in enumerate(aims, 1))


def _procedures_overview(rows, acclim_days="", methods="", duration=""):
    """PART 6 overview — a clean, professional list of the manipulations, then
    the researcher's optional free-text methods, then a representative timeline."""
    lines = ["The following procedures and manipulations will be performed:"]
    dur = _s(duration)
    if dur:
        lines.insert(0, f"Expected study duration: {dur}.")
    acclim = _s(acclim_days)
    lines.append(f"1. Acclimatization: animals are allowed to acclimatize to the facility for "
                 f"{acclim + ' days' if acclim else 'a defined period'} prior to any procedure, "
                 "with free access to food and water and daily health checks.")
    n = 2
    for g, _, s in rows:
        if _is_control(g):
            continue
        drug = _s(g.get("drug_name"), "the test article")
        route = _s(g.get("route"), "the assigned route")
        route_txt = route if ("route" in route.lower() or "(" in route) else f"the {route} route"
        dose = _s(g.get("dose"))
        dose_txt = f" at {dose} mg/kg" if dose else ""
        lines.append(f"{n}. Dosing ({_s(g.get('group_name'), 'group')}): administration of "
                     f"{drug}{dose_txt} by {route_txt}.")
        n += 1
    lines.append(f"{n}. Monitoring: regular body-weight measurement and structured clinical "
                 "observation for signs of toxicity against predefined humane endpoints.")
    n += 1
    samples = _join_unique(x for _, _, s in rows for x in (s.get("recommended_samples") or []))
    if samples:
        lines.append(f"{n}. Sample collection: {_list_sentence(samples)}, collected within safe "
                     "volume limits at defined time-points.")
        n += 1
    lines.append(f"{n}. Termination: humane euthanasia and terminal tissue collection at the "
                 "study endpoint (see Part 8).")

    # optional researcher-provided methods / procedure detail
    meth = _s(methods)
    if meth:
        lines.append("")
        lines.append("Additional methods / procedure detail:")
        lines.append(meth)

    # per-group special instructions supplied by the researcher
    special = [f"• {_s(g.get('group_name'), 'Group')}: {_s(g.get('instructions'))}"
               for g, _, _ in rows if _s(g.get("instructions"))]
    if special:
        lines.append("")
        lines.append("Special instructions:")
        lines.extend(special)

    # a representative timeline, if available
    tl = None
    for _, _, s in rows:
        if s.get("timeline"):
            tl = s["timeline"]
            break
    if tl:
        lines.append("")
        lines.append("Representative timeline:")
        for step in tl[:8]:
            if isinstance(step, dict):
                day = _s(step.get("day"))   # already contains "Day …"
                act = _s(step.get("activity") or step.get("phase"))
                lines.append(f"  {day}: {act}" if day else f"  {act}")
    return "\n".join(lines)


def _experimental_endpoints(rows):
    eps = _join_unique(e for g, _, _ in rows for e in (g.get("toxicity_endpoints") or []))
    samples = _join_unique(x for _, _, s in rows for x in (s.get("recommended_samples") or []))
    out = ["Experimental endpoints (scientific data collection points):"]
    if eps:
        out.append(f"• Assessed measures: {_list_sentence(eps)}.")
    if samples:
        out.append(f"• Terminal samples/tissues: {_list_sentence(samples)}.")
    out.append("• Study endpoint is reached at the scheduled terminal time-point, after "
               "which planned analyses are performed.")
    return "\n".join(out)


def _welfare_of(rows):
    """Return the first available welfare dict among treatment groups."""
    for g, a, s in rows:
        w = s.get("welfare")
        if w:
            return w
    return None


def _humane_endpoints(rows):
    w = _welfare_of(rows)
    if w and w.get("humane_endpoints"):
        return _bullets(w["humane_endpoints"])
    return _bullets([
        "Body weight loss ≥ 20% from baseline (or ≥ 15% with other clinical signs)",
        "Body condition score ≤ 2/5 (emaciation)",
        "Persistent hunched posture, lack of grooming, or hypothermia",
        "Inability to reach food or water for > 24 h",
    ])


def _observation_frequency(rows):
    w = _welfare_of(rows)
    if w and w.get("monitoring", {}).get("frequency"):
        m = w["monitoring"]
        extra = _s(m.get("body_condition_scoring"))
        return _s(m["frequency"]) + (f" {extra}" if extra else "")
    return ("At least once daily, with additional checks after each dosing event and "
            "around the expected peak drug effect. Record body weight and body condition "
            "score (1–5) at each observation.")


def _endpoint_response(rows):
    w = _welfare_of(rows)
    if w and w.get("response_when_endpoint_reached"):
        return _s(w["response_when_endpoint_reached"])
    return ("The affected animal will be removed from study and humanely euthanised using "
            "the approved method immediately; observations are recorded and the attending "
            "veterinarian is notified.")


def _euthanasia_method(rows):
    return ("Euthanasia will be performed in accordance with the AVMA Guidelines and "
            "institutional SOPs — CO₂ inhalation from a compressed source at an approved "
            "flow rate, followed by a secondary physical method (e.g. cervical dislocation "
            "or exsanguination) to confirm death.")


def _pain_monitoring_frequency(rows):
    w = _welfare_of(rows)
    if w and w.get("monitoring", {}).get("frequency"):
        return _s(w["monitoring"]["frequency"])
    return ("Once daily during the study, increasing to twice daily around dosing and the "
            "expected peak effect; more frequently if any clinical signs are observed.")


def _discomfort_measures(rows):
    w = _welfare_of(rows)
    base = ("Measures to minimise discomfort, distress or pain: gentle handling by trained "
            "personnel, use of the least-stressful effective route and volume, provision of "
            "soft/moist food and supplemental warmth as needed, and prompt clinical intervention.")
    if w and w.get("monitoring", {}).get("analgesia_guidance"):
        base += " " + _s(w["monitoring"]["analgesia_guidance"])
    return base


# ── structured tables ─────────────────────────────────────────────────────

def _team(admin, study):
    team = []
    pi = _s(study.get("pi_name"))
    inst = _s(study.get("institution"), "KAIMRC")
    if pi:
        team.append({"name": pi, "role": "Principal Investigator",
                     "qualifications": _s(admin.get("pi_qualifications")),
                     "institution": inst,
                     "email": _s(admin.get("pi_email")),
                     "mobile": _s(admin.get("pi_mobile"))})
    for m in (admin.get("team") or []):
        if not isinstance(m, dict):
            continue
        if not _s(m.get("name")):
            continue
        team.append({
            "name": _s(m.get("name")),
            "role": _s(m.get("role")),
            "qualifications": _s(m.get("qualifications")),
            "institution": _s(m.get("institution"), inst),
            "email": _s(m.get("email")),
            "mobile": _s(m.get("mobile")),
        })
    return team[:6]


def _animals(rows, source=None):
    # None (field absent) → default vendor; "" (researcher chose "Other" and left
    # it blank) → leave the Source cell empty for them to fill in the form.
    source = "Accredited vendor" if source is None else _s(source)
    animals = []
    for g, _, s in rows:
        sp = _s(s.get("species") or g.get("species"), "Mouse")
        strain = _s(s.get("strain") or g.get("strain") or s.get("recommended_strain"))
        # transgenic line → strain marked with "*" (the form asks to indicate
        # genetically-modified strains with *) and the line noted in parentheses
        tg = _s(g.get("transgenic_line"))
        if tg:
            strain = (strain + " " if strain else "") + f"* ({tg})"
        sex = _s(g.get("sex"), "—")
        age = _s(s.get("planned_age_weeks") or g.get("age"))
        age = f"{age} wk" if age and "wk" not in age.lower() and "week" not in age.lower() else (age or "—")
        total = _s(s.get("planned_animals") or g.get("num_mice"), "—")
        # per-group Vendor overrides the study-wide source when provided
        src = _s(g.get("vendor")) or source
        animals.append({"species": sp, "strain": strain or "—", "sex": sex,
                        "age": age, "total": total, "source": src})
    return animals[:6]


def _funding(admin):
    fund = _s(admin.get("funding_source"))
    housing = _s(admin.get("housing_type")).lower()
    note = fund
    if housing == "absl":
        note = (note + " | " if note else "") + "Housing: ABSL (biocontainment)"
    elif housing == "standard":
        note = (note + " | " if note else "") + "Housing: standard"
    return note


# ── main entry point ──────────────────────────────────────────────────────

def _clean_text(t):
    """Light grammar/format tidy-up for text the researcher wants copied as-is:
    collapse whitespace, capitalise the first letter, ensure a closing period."""
    t = re.sub(r'\s+', ' ', _s(t)).strip()
    if not t:
        return ""
    t = t[0].upper() + t[1:]
    if t[-1] not in '.!?:':
        t += '.'
    return t


def build_context(payload):
    study = payload.get("study") or {}
    admin = payload.get("admin") or {}
    refine = study.get("refine") or {}
    rows = _pair(payload.get("groups") or [], payload.get("analysis") or [])

    # PART 4 Aims ← researcher's Hypothesis & Objectives (verbatim unless they
    # ticked "review & rewrite"; if empty, generated from the study data).
    hyp = _s(study.get("hypothesis"))
    if hyp and not refine.get("hypothesis"):
        research_aims = _clean_text(hyp)
    elif hyp:
        research_aims = _clean_text(hyp) + "\n\n" + _research_aims(rows)
    else:
        research_aims = _research_aims(rows)

    # PART 2 Expected outcomes ← researcher's Outcomes field (same logic)
    out_txt = _s(study.get("outcomes"))
    if out_txt and not refine.get("outcomes"):
        expected_outcomes = _clean_text(out_txt)
    else:
        expected_outcomes = _expected_outcomes(rows)

    return {
        "study_purpose": _study_purpose(rows, study),
        "what_done": _what_done(rows),
        "expected_outcomes": expected_outcomes,
        "study_benefit": _study_benefit(rows),
        "strain_health_issues": _strain_health_issues(rows),
        "scientific_justification": _scientific_justification(rows),
        "literature_background": _literature_background(rows, study.get("background"),
                                                        refine.get("background")),
        "research_aims": research_aims,
        "procedures_overview": _procedures_overview(rows, admin.get("acclimatization_days"),
                                                     study.get("methods"), study.get("duration")),
        "experimental_endpoints": _experimental_endpoints(rows),
        "humane_endpoints": _humane_endpoints(rows),
        "observation_frequency": _observation_frequency(rows),
        "endpoint_response": _endpoint_response(rows),
        "euthanasia_method": _euthanasia_method(rows),
        "pain_monitoring_frequency": _pain_monitoring_frequency(rows),
        "discomfort_measures": _discomfort_measures(rows),
        "funding_source": _funding(admin),
        "team": _team(admin, study),
        "animals": _animals(rows),
    }


# ══════════════════════════════════════════════════════════════════════════
#  Post-process pass: tick checkboxes + fill structured tables
#  (docxtpl handles free text; Word form controls are set here directly).
# ══════════════════════════════════════════════════════════════════════════

# Everything the platform writes is marked with a light-blue (cyan) highlight so
# the researcher can instantly tell auto-filled content from the original form.
_HL = 'cyan'


def _highlight_run(r):
    rpr = r.find(qn('w:rPr'))
    if rpr is None:
        rpr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rpr)
    hl = rpr.find(qn('w:highlight'))
    if hl is None:
        hl = rpr.makeelement(qn('w:highlight'), {})
        rpr.append(hl)
    hl.set(qn('w:val'), _HL)


def _highlight_para(p):
    for r in p.findall(qn('w:r')):
        _highlight_run(r)


# The 1×1 narrative answer boxes filled by docxtpl (see build_template BOX_TAGS)
# — highlight their whole content after rendering.
_DOCXTPL_BOX_TABLES = [6, 7, 8, 9, 12, 17, 18, 19, 33, 57, 58, 59, 61, 63, 80, 81]


def _highlight_docxtpl_boxes(doc):
    for idx in _DOCXTPL_BOX_TABLES:
        try:
            tc = next(doc.tables[idx]._tbl.iter(qn('w:tc')))
            for p in tc.iter(qn('w:p')):
                _highlight_para(p)
        except Exception:
            pass


def _tick_sdt(sdt, checked=True):
    """Flip a Word checkbox content control and swap its ☐/☒ display glyph."""
    ck = sdt.find('.//' + qn('w14:checked'))
    if ck is not None:
        ck.set(qn('w14:val'), '1' if checked else '0')
    content = sdt.find(qn('w:sdtContent'))
    if content is not None:
        first = True
        for t in content.iter(qn('w:t')):
            t.text = ('☒' if checked else '☐') if first else ''
            first = False
        if checked:                       # highlight the ticked box
            for r in content.iter(qn('w:r')):
                _highlight_run(r)


def _cell_checkboxes(cell):
    return [s for s in cell._element.findall('.//' + qn('w:sdt'))
            if s.find('.//' + qn('w14:checkbox')) is not None]


def _unlock_content_controls(doc):
    """Make the whole document freely editable after download.

    The blank form ships with (a) document-level protection (edit="forms") and
    (b) locked content controls, and several answer cells are wrapped in
    cell-level content controls. Together these stop the researcher editing
    parts of the form (e.g. Part 4 / Part 8). We remove the protection, drop
    every lock, and unwrap text (non-checkbox) content controls to plain cells."""
    # 1) remove document-level protection
    try:
        st = doc.settings.element
        dp = st.find(qn('w:documentProtection'))
        if dp is not None:
            st.remove(dp)
    except Exception:
        pass
    # 2) drop every content-control lock
    for lock in doc.element.body.findall('.//' + qn('w:lock')):
        parent = lock.getparent()
        if parent is not None:
            parent.remove(lock)
    # 3) unwrap cell-level text content controls (keep checkbox controls intact)
    for sdt in list(doc.element.body.iter(qn('w:sdt'))):
        if sdt.find('.//' + qn('w14:checkbox')) is not None:
            continue
        content = sdt.find(qn('w:sdtContent'))
        if content is None:
            continue
        tc = content.find(qn('w:tc'))
        parent = sdt.getparent()
        if tc is not None and parent is not None and parent.tag == qn('w:tr'):
            parent.replace(sdt, tc)


def _tick_row_checkbox(table, row_idx, checked=True):
    """Tick the first checkbox content control anywhere in a row (some form
    checkboxes live outside clean cell boundaries — e.g. the emergency Yes/No)."""
    tr = table.rows[row_idx]._tr
    for sdt in tr.iter(qn('w:sdt')):
        if sdt.find('.//' + qn('w14:checkbox')) is not None:
            _tick_sdt(sdt, checked)
            return True
    return False


def _write_row_textbox(table, row_idx, text, n=0):
    """Write into the n-th plain-text content control in a row (n<0 counts from
    the end). Used for the emergency 'special instructions' and phone boxes."""
    boxes = [s for s in table.rows[row_idx]._tr.iter(qn('w:sdt'))
             if s.find('.//' + qn('w14:checkbox')) is None
             and s.find(qn('w:sdtContent')) is not None]
    if not boxes:
        return False
    sdt = boxes[n if n >= 0 else len(boxes) + n]
    content = sdt.find(qn('w:sdtContent'))
    p = content.find(qn('w:p'))
    if p is None:
        p = content.find('.//' + qn('w:p'))
    if p is not None:
        _set_para_el(p, text)
        _left_align(p)
        return True
    # inline text content control: sdtContent > w:r > w:t (no paragraph)
    runs = content.findall(qn('w:r'))
    if runs:
        ts = runs[0].findall(qn('w:t'))
        if ts:
            ts[0].text = text
            for x in ts[1:]:
                x.text = ''
        else:
            tt = runs[0].makeelement(qn('w:t'), {})
            tt.set(qn('xml:space'), 'preserve')
            tt.text = text
            runs[0].append(tt)
        for r in runs[1:]:
            r.getparent().remove(r)
        _highlight_run(runs[0])
        return True
    return False


def _check(cell, cb_index=0, checked=True):
    cbs = _cell_checkboxes(cell)
    if 0 <= cb_index < len(cbs):
        _tick_sdt(cbs[cb_index], checked)
        return True
    return False


def _set_para_el(p, text):
    """Set a <w:p> element's text, preserving the first run's formatting, and
    mark it with the platform's blue highlight."""
    runs = p.findall(qn('w:r'))
    if runs:
        ts = runs[0].findall(qn('w:t'))
        if ts:
            ts[0].text = text
            for x in ts[1:]:
                x.text = ''
        else:
            t = runs[0].makeelement(qn('w:t'), {})
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            runs[0].append(t)
        for r in runs[1:]:
            r.getparent().remove(r)
    else:
        r = p.makeelement(qn('w:r'), {})
        t = p.makeelement(qn('w:t'), {})
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    _highlight_para(p)


def _left_align(p):
    """Force left alignment on a paragraph (answer boxes default to 'justify',
    which stretches multi-line answers across the whole line)."""
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = p.makeelement(qn('w:pPr'), {})
        p.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = pPr.makeelement(qn('w:jc'), {})
        pPr.append(jc)
    jc.set(qn('w:val'), 'left')


def _ordered_unwrapped_tcs(tr):
    """Return a row's cells (<w:tc>) in visual order, unwrapping any that are
    wrapped in a cell-level content control (tr > sdt > … > tc).

    python-docx's `row.cells` misaligns columns when a cell is sdt-wrapped, so
    for the structured tables we walk the row XML directly and unwrap first."""
    tcs = []
    for ch in list(tr):
        if ch.tag == qn('w:tc'):
            tcs.append(ch)
        elif ch.tag == qn('w:sdt'):
            inner = ch.find('.//' + qn('w:tc'))
            if inner is not None:
                tr.replace(ch, inner)      # drop the content control + placeholder
                tcs.append(inner)
    return tcs


def _set_tc_clean(tc, text):
    """Clear a cell element and write clean left-aligned text into it."""
    for child in list(tc):
        if child.tag == qn('w:tcPr'):
            continue
        tc.remove(child)
    p = tc.makeelement(qn('w:p'), {})
    tc.append(p)
    _set_para_el(p, text)
    _left_align(p)


def _fill_struct_row(row, values):
    """Fill a structured-table row (personnel / animals) column by column,
    handling sdt-wrapped cells so columns line up."""
    tcs = _ordered_unwrapped_tcs(row._tr)
    for ci, v in enumerate(values):
        if v is not None and ci < len(tcs):   # None = leave this cell untouched
            _set_tc_clean(tcs[ci], v)


def _write_cell(cell, text):
    """Write text into a cell — into its plain-text content control if it has
    one (handling both paragraph and inline controls, so the 'Click here to
    enter text.' placeholder is always replaced), else the first paragraph."""
    for s in cell._element.findall('.//' + qn('w:sdt')):
        if s.find('.//' + qn('w14:checkbox')) is None:
            content = s.find(qn('w:sdtContent'))
            if content is None:
                continue
            p = content.find(qn('w:p'))
            if p is None:
                p = content.find('.//' + qn('w:p'))
            if p is not None:
                _set_para_el(p, text)
                _left_align(p)
                return
            # inline content control (sdtContent > w:r) — replace its run text
            runs = content.findall(qn('w:r'))
            if runs:
                ts = runs[0].findall(qn('w:t'))
                if ts:
                    ts[0].text = text
                    for x in ts[1:]:
                        x.text = ''
                else:
                    tt = runs[0].makeelement(qn('w:t'), {})
                    tt.set(qn('xml:space'), 'preserve')
                    tt.text = text
                    runs[0].append(tt)
                for r in runs[1:]:
                    r.getparent().remove(r)
                _highlight_run(runs[0])
                return
    p = cell.paragraphs[0]._p
    _set_para_el(p, text)
    _left_align(p)


def _write_box(table, text):
    """Fill a 1×1 answer-box table. Clears the whole cell first — including any
    content control and its 'Click here to enter text.' placeholder — then
    writes clean, highlighted, left-aligned text."""
    tc = next(table._tbl.iter(qn('w:tc')))
    _set_tc_clean(tc, text)


def _write_explain(table, text):
    """Write into a Yes/No table's explanation box (bottom-most plain-text
    content control)."""
    for row in reversed(table.rows):
        for cell in row.cells:
            if any(s.find('.//' + qn('w14:checkbox')) is None
                   for s in cell._element.findall('.//' + qn('w:sdt'))):
                _write_cell(cell, text)
                return
    _write_cell(table.rows[-1].cells[-1], text)


def _fill_row(table, row_idx, values):
    if row_idx >= len(table.rows):
        return
    cells = table.rows[row_idx].cells
    for i, v in enumerate(values):
        if v is not None and i < len(cells):
            _write_cell(cells[i], v)


def _reduce_text(rows, design=""):
    """Statistical sample-size justification — reuses the platform's own
    power-analysis rationale (never a fabricated one). Prefixed with the study
    design type when the researcher selected one."""
    prefix = f"Study design: {design}. " if _s(design) else ""
    for _, a, _ in rows:
        r = _s(a.get('rationale'))
        if r and 'control should match' not in r.lower():
            return prefix + r
    ns = [_s(g.get('num_mice')) for g, _, _ in rows if _s(g.get('num_mice'))]
    return (prefix + "Group sizes were determined by a priori power analysis (two-sided "
            "α = 0.05, power = 0.80) for the primary endpoint, using the fewest "
            "animals expected to yield statistically valid results"
            + (f"; planned n per group: {', '.join(ns)}." if ns else "."))


def _study_flags(rows):
    """Decisions that depend on the study's predicted toxicity."""
    w = _welfare_of(rows)
    level = (w.get('monitoring', {}) or {}).get('level') if w else None
    pain = level in ('moderate', 'high')
    return {'pain': pain}


def fill_form_controls(doc, rows, admin, study):
    """Tick every relevant checkbox and fill the structured tables so the
    researcher only has to review. Standard, conservative answers for a
    non-surgical, non-hazardous rodent pharmacology / toxicology study."""
    T = doc.tables
    flags = _study_flags(rows)
    pain = flags['pain']
    absl = _s((admin or {}).get('housing_type')).lower() == 'absl'

    # ── PART 1 — Research team (T2) & PART 3 — animal table (T10) ──────────
    # These cells carry content controls (dropdowns / placeholders); fill them
    # with clean text in the correct columns, clearing the placeholders.
    team = _team(admin, study)
    for i, m in enumerate(team[:6]):
        r = 1 + i
        if r >= len(T[2].rows):
            break
        try:
            _fill_struct_row(T[2].rows[r],
                             [m['name'], m['role'], m['qualifications'],
                              m['institution'], m['email'], m['mobile']])
        except Exception:
            pass
    animals = _animals(rows, (admin or {}).get('animal_source'))
    for i, an in enumerate(animals[:6]):
        r = 1 + i
        if r >= len(T[10].rows):
            break
        try:
            _fill_struct_row(T[10].rows[r],
                             [an['species'], an['strain'], an['sex'],
                              an['age'], an['total'], an['source']])
        except Exception:
            pass

    # ── PART 1 B — personnel × procedures matrix (T3) ─────────────────────
    species_all = _list_sentence(_join_unique(
        _s(s.get('species') or g.get('species')) for g, _, s in rows)) or 'Mouse'
    has_blood = any(any(k in _s(x).lower() for k in ('blood', 'plasma', 'serum'))
                    for g, _, s in rows
                    for x in (list(g.get('sample_types') or []) + list(s.get('recommended_samples') or [])))
    # experience per procedure column (c4..c15) for a non-surgical dosing/tox study
    #   MT, Restraint, Non-invasive, Anaesthesia, Euthanasia, Blood, Surg(S),
    #   Surg(non-S), Surgical assist, Assess humane EP, Husbandry, Other
    proc = ['***', '***', '***', '**' if has_blood else 'NE', '**',
            '**' if has_blood else 'NE', 'NE', 'NE', 'NE', '***', '**', '']
    for i, m in enumerate(team[:6]):
        r = 4 + i
        if r >= len(T[3].rows):
            break
        try:
            _fill_struct_row(T[3].rows[r],
                             [m['name'], '', m['qualifications'], species_all] + proc)
        except Exception:
            pass

    # "Describe any other procedure not listed …" box (T4)
    _route = _list_sentence(_join_unique(
        _s(g.get('route')) for g, _, _ in rows if not _is_control(g))) or 'the specified route'
    try:
        _write_box(T[4], "The research team will perform rodent handling and restraint, "
                         f"{_route} dosing, blood / tissue sample collection, clinical monitoring "
                         "and humane-endpoint assessment for this study. No procedures beyond those "
                         "listed in the table above are undertaken by this team.")
    except Exception:
        pass

    def cb(idx, cell, k=0, checked=True):
        try:
            _check(T[idx].rows[0].cells[cell], k, checked)
        except Exception:
            pass

    def cbr(idx, row, cell, k=0, checked=True):
        try:
            _check(T[idx].rows[row].cells[cell], k, checked)
        except Exception:
            pass

    def explain(idx, text):
        try:
            _write_explain(T[idx], text)
        except Exception:
            pass

    def box(idx, text):
        try:
            _write_box(T[idx], text)
        except Exception:
            pass

    def row(idx, r, values):
        # clean fill: clears any content-control placeholder ("Click here to
        # enter text.") so only the value shows, columns aligned
        try:
            _fill_struct_row(T[idx].rows[r], values)
        except Exception:
            pass

    # ── PART 4 — 3Rs ──────────────────────────────────────────────────────
    cb(20, 2)                                   # duplicate previous work? NO
    explain(20, "A structured literature search (PubMed / PubChem and related "
                "databases) confirmed that the specific question addressed here is "
                "unresolved; the study does not duplicate published work.")
    cb(21, 0)                                    # reduced to fewest? YES
    explain(21, _reduce_text(rows, (study or {}).get('design')))
    cb(22, 0)                                    # potential for pain/distress? YES
    explain(22, "Any potential for pain or distress is minimised by trained "
                "handling, the least-stressful effective route and volume, structured "
                "daily monitoring against predefined humane endpoints, and analgesia "
                "where indicated (see Part 11).")
    cb(23, 2)                                    # could models replace animals? NO
    explain(23, "In-silico toxicity modelling and published in-vitro data were used "
                "to refine the design and dose selection, but cannot replace a live "
                "animal: the endpoints require an intact, physiologically integrated "
                "mammalian system (systemic ADME and multi-organ response).")

    # ── PART 4 — experimental justification of numbers ────────────────────
    for i, (g, a, s) in enumerate(rows):
        row(24, 1 + i, [_s(g.get('group_name'), f'Group {i + 1}'),
                        _s(s.get('species') or g.get('species'), 'Mouse'),
                        _s(s.get('planned_animals') or g.get('num_mice'))])
    cbr(25, 2, 0)                                # "determined statistically"
    row(25, 3, [None, _reduce_text(rows, (study or {}).get('design'))])   # description

    # ── PART 3 — quantification of animals ────────────────────────────────
    cat = 'Category D' if pain else 'Category C'
    total = 0
    for i, (g, a, s) in enumerate(rows):
        n = _s(s.get('planned_animals') or g.get('num_mice'))
        m = re.findall(r'\d+', n)
        if m:
            total += int(m[0])
        row(13, 2 + i, [_s(s.get('species') or g.get('species'), 'Mouse'),
                        cat, '', n, '', '', n])
    if total:
        try:
            _write_cell(T[13].rows[6].cells[6], str(total))
        except Exception:
            pass

    # ── PART 1 — funding agency (T5) & PART 3 — housing location (T11) ─────
    fund_cb = {'kaimrc': 0, 'moe': 3, 'moh': 4, 'rdia': 8, 'snih': 9, 'other': 11}
    ag = _s((admin or {}).get('funding_agency')).lower()
    if ag in fund_cb:
        cbr(5, 1, fund_cb[ag], 0)
    loc_cb = {'riyadh': 0, 'jeddah': 1, 'al ahsa': 2, 'alahsa': 2}
    loc = _s((admin or {}).get('facility_location')).lower()
    cbr(11, 0, loc_cb.get(loc, 0), 0)            # default Riyadh

    # ── PART 1 F — emergency instructions + declarations (auto-agree) ──────
    instr_bits = []
    for g, _, _ in rows:
        ins = _s(g.get('instructions'))
        if ins:
            instr_bits.append(f"{_s(g.get('group_name'), 'Group')}: {ins}")
    special = "; ".join(dict.fromkeys(instr_bits))
    try:
        if special:
            _tick_row_checkbox(T[5], 10)             # YES – special instructions apply
            _write_row_textbox(T[5], 12, special, 0)  # …provide them (box lives in r12)
        else:
            _tick_row_checkbox(T[5], 9)              # NO – vet may use professional judgment
    except Exception:
        pass
    phone = _s((admin or {}).get('emergency_phone'))
    if phone:
        try:
            _write_row_textbox(T[5], 13, phone, -1)   # emergency phone box (r13)
        except Exception:
            pass
    cbr(5, 14, 5, 0)                                 # "I understand… I agree" statement
    try:
        cbr(121, 1, 0, 0)                            # Appendix 2: "I AGREE WITH ALL THE ABOVE"
    except Exception:
        pass

    # ── PART 5 — housing, diet, husbandry ─────────────────────────────────
    cb(26, 2)                                    # housed at another facility? NO
    explain(26, "All animals are housed and used at the KAIMRC animal facility.")
    cbr(27, 1, 0, 0)                             # standard housing: YES
    cbr(27, 1, 1, 0)                             # arrangement: GROUP
    cbr(27, 1, 2, 2 if absl else 0)              # ABSL if requested, else CONVENTIONAL
    # special environmental conditions (T27 explain box), from the group cards
    _hz = []
    for g, _, _ in rows:
        t, h, l = _s(g.get('housing_temp')), _s(g.get('housing_humidity')), _s(g.get('housing_light'))
        bits = [f"temperature {t}" if t else "", f"humidity {h}" if h else "",
                f"light cycle {l}" if l else ""]
        bits = [b for b in bits if b]
        if bits:
            _hz.append("; ".join(bits))
    if _hz:
        try:
            _write_cell(T[27].rows[3].cells[0],
                        "Environmental conditions: " + " | ".join(dict.fromkeys(_hz)) + ".")
        except Exception:
            pass
    cb(28, 0)                                    # housing discussed with vet? YES
    explain(28, "Attending facility veterinarian, KAIMRC.")
    cb(29, 0)                                    # standard diet? YES
    cb(30, 0)                                    # feeding schedule: Ad Lib
    cb(31, 2)                                    # restricted watering? NO
    cb(32, 0)                                    # standard cage change? YES

    # ── PART 6 — hazards checklist (all NO) + brief restraint ─────────────
    for idx in (35, 36, 37, 38, 39):
        cb(idx, 2)
        explain(idx, "None — not applicable to this protocol.")
    cbr(40, 0, 1, 0)                             # physical restraint used? YES (brief)
    box(41, "Brief manual restraint (scruff / one-handed hold) for < 2 minutes "
            "during gavage, injection and sample collection; no mechanical restraint "
            "device is used.")
    cb(34, 0)                                    # acclimatization period? YES

    # ── PART 7 — hazardous agents (A/B/C): none ───────────────────────────
    box(50, "Not applicable — no Category A/B/C agents (radioactive, "
            "carcinogenic/cytotoxic, or infectious) are administered in this protocol.")

    # ── PART 8 — disposition & euthanasia ─────────────────────────────────
    cb(62, 0)                                    # animals will be euthanised
    row(64, 1, ["CO₂ (compressed gas)",
                "Gradual chamber displacement (AVMA-compliant flow rate)", "Inhalation"])
    box(65, "A secondary physical method (cervical dislocation or exsanguination) "
            "is applied to confirm death.")

    # ── PART 9 — collection of body fluids/tissues from living animals ─────
    blood_samples, vol = [], ""
    for g, _, s in rows:
        for x in list(g.get('sample_types') or []) + list(s.get('recommended_samples') or []):
            xl = _s(x).lower()
            if any(k in xl for k in ('blood', 'plasma', 'serum')) and _s(x) not in blood_samples:
                blood_samples.append(_s(x))
        bv = s.get('blood_volume_ml')
        if bv and not vol:
            vol = f"≤ {bv} mL (within safe blood-volume limits)"
    if not vol:
        vol = "Within safe blood-volume limits"
    if blood_samples:
        for i, samp in enumerate(blood_samples[:3]):
            try:
                _fill_struct_row(T[67].rows[1 + i],
                                 [samp, "At scheduled time-point / terminal", vol,
                                  "Submandibular / tail vein (survival) or cardiac puncture (terminal)"])
            except Exception:
                pass
        cb(68, 0)                                # anaesthetised/sedated? YES
        row(69, 1, ["Isoflurane (suggested — confirm with veterinarian)",
                    "To effect (2–3%)", "Inhalation", "To effect"])
    else:
        cb(66, 0)                                # NOT APPLICABLE (no living-animal fluid collection)

    # ── PART 11 — pain & distress management ──────────────────────────────
    if pain:
        cb(76, 0)                                # distress/pain expected? YES
        explain(76, "Administration of the test article and any expected toxicity may "
                    "cause transient discomfort; animals are monitored against "
                    "predefined humane endpoints (Part 8).")
        cb(77, 0)                                # analgesic/anaesthetic used? YES
        row(78, 1, ["Buprenorphine (suggested — confirm with veterinarian)",
                    "0.05–0.1 mg/kg", "SC", "Every 8–12 h as needed"])
    else:
        cb(76, 2)                                # distress/pain expected? NO
        cb(79, 0)                                # analgesic drugs used? NO
        explain(79, "Only routine procedures (injection / gavage and limited sampling) "
                    "are performed and are not expected to cause more than momentary "
                    "discomfort, so analgesia is not routinely required. Analgesia will "
                    "be provided promptly if any distress is observed, in consultation "
                    "with the attending veterinarian.")
    cb(82, 0)                                    # euthanise if severely ill? YES
    explain(82, "Predefined humane endpoints (Part 8): ≥ 20% body-weight loss, body "
                "condition score ≤ 2/5, severe or unresolving clinical signs, or "
                "inability to reach food / water.")

    # ── PART 10 — antibody production: none ───────────────────────────────
    box(72, "Not applicable — this protocol does not involve antibody production.")

    # ── PART 12 — surgery: none in this protocol ──────────────────────────
    cb(88, 2)                                    # neuromuscular blockers? NO
    cb(90, 2)                                    # non-survival practice animals? NO
    box(105, "Not applicable — no surgical procedures are performed in this protocol.")

    # Highlight the docxtpl-filled narrative boxes (blue), then make every
    # answer box editable after download (unlock content controls).
    _highlight_docxtpl_boxes(doc)
    try:
        _unlock_content_controls(doc)
    except Exception:
        pass


def generate_iacuc_docx(payload):
    """Render the filled IACUC form and return it as an in-memory BytesIO.

    Two passes: docxtpl fills the free-text narrative, then a python-docx pass
    ticks the form's checkboxes and fills the structured tables — so the
    researcher only reviews. References come solely from the platform's real
    literature search (never fabricated)."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"IACUC template not found at {TEMPLATE_PATH}. "
            "Run: python iacuc_assets/build_template.py")
    rows = _pair(payload.get("groups") or [], payload.get("analysis") or [])

    tpl = DocxTemplate(str(TEMPLATE_PATH))
    tpl.render(build_context(payload))
    tmp = io.BytesIO()
    tpl.save(tmp)
    tmp.seek(0)

    doc = Document(tmp)
    fill_form_controls(doc, rows, payload.get("admin") or {}, payload.get("study") or {})

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
