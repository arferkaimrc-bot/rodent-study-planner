"""
Enhanced Experimental Rodent Study Planner - COMPLETE VERSION
=============================================================
Features:
- Multiple scientific paper APIs (NCBI, Semantic Scholar, OpenAlex, CrossRef)
- Machine Learning for sample size prediction
- Blood quantity calculator
- Enhanced PDF/Word exports with diagrams and colors
- Power analysis curves
- Study timeline visualizations

Author: Enhanced Platform v3.0
Version: 3.0
"""

from flask import Flask, request, jsonify, render_template, send_file
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import requests
from io import BytesIO
import logging
from datetime import datetime, timedelta
import os
import math
import re
import hashlib
import pickle
import numpy as np
import pandas as pd
from functools import lru_cache
from tenacity import retry, stop_after_attempt, wait_exponential
from markupsafe import escape
import time
import json

# PDF / Word generation
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image, KeepTogether
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.graphics.shapes import Drawing, Rect, String, Line
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from iacuc_generator import generate_iacuc_docx

# Statistical packages
from scipy import stats
from statsmodels.stats.power import tt_ind_solve_power

# ML packages
try:
    from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
    from sklearn.preprocessing import StandardScaler
    from sklearn.model_selection import train_test_split, cross_val_score
    import joblib
    ML_AVAILABLE = True
except ImportError:   
    ML_AVAILABLE = False
    print("Warning: scikit-learn not installed. ML features disabled.")

# Validation
try:
    from marshmallow import Schema, fields, validate, ValidationError, EXCLUDE
    VALIDATION_AVAILABLE = True
except ImportError:
    VALIDATION_AVAILABLE = False
    print("Warning: marshmallow not installed. Advanced validation disabled.")

# RDKit for molecular descriptors (REQUIRED for toxicity prediction)
try:
    from rdkit import Chem
    from rdkit.Chem import Descriptors, Crippen, Lipinski, rdMolDescriptors
    RDKIT_AVAILABLE = True
except ImportError:
    RDKIT_AVAILABLE = False
    logger_msg = "RDKit not installed. Toxicity prediction will use simplified model."
    print(f"Warning: {logger_msg}")

# ============================================================================
# CONFIGURATION
# ============================================================================

class Config:
    """Application configuration."""
    DEBUG = os.getenv('FLASK_DEBUG', 'False').lower() == 'true'
    SECRET_KEY = os.getenv('SECRET_KEY', 'dev-secret-key-change-in-production')
    
    # API Configuration
    NCBI_API_KEY = os.getenv('NCBI_API_KEY', None)
    NCBI_API_BASE = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    PUBCHEM_API_BASE = "https://pubchem.ncbi.nlm.nih.gov/rest/pug"
    EUROPE_PMC_BASE = "https://www.ebi.ac.uk/europepmc/webservices/rest"
    IMPC_BASE = "https://www.ebi.ac.uk/mi/impc/solr"
    
    # NEW APIS
    SEMANTIC_SCHOLAR_BASE = "https://api.semanticscholar.org/graph/v1"
    OPENALEX_BASE = "https://api.openalex.org"
    CROSSREF_BASE = "https://api.crossref.org/works"
    DOAJ_BASE = "https://doaj.org/api/v2/search/articles"   # open-access journals
    
    # ChEMBL API for bioactivity data
    CHEMBL_BASE = "https://www.ebi.ac.uk/chembl/api/data"
    CHEMBL_TIMEOUT = 15

    # PubChem PUG-View + SDQ (real experimental acute-toxicity / LD50 data from ChemIDplus)
    PUBCHEM_PUGVIEW_BASE = "https://pubchem.ncbi.nlm.nih.gov/rest/pug_view/data/compound"
    PUBCHEM_SDQ = "https://pubchem.ncbi.nlm.nih.gov/sdq/sdqagent.cgi"

    # Trained LD50 regression model (real data: TDC LD50_Zhu, rat oral)
    LD50_MODEL_PATH = './ml_models/ld50_model.pkl'
    LD50_META_PATH = './ml_models/ld50_meta.json'

    # Trained aqueous-solubility model (real data: TDC Solubility_AqSolDB)
    SOLUBILITY_MODEL_PATH = './ml_models/solubility_model.pkl'
    SOLUBILITY_META_PATH = './ml_models/solubility_meta.json'

    # Trained half-life model (real data: TDC Half_Life_Obach — HUMAN PK, relative)
    HALFLIFE_MODEL_PATH = './ml_models/halflife_model.pkl'
    HALFLIFE_META_PATH = './ml_models/halflife_meta.json'

    # Batch of binary safety/ADME classifiers (loaded generically from ml_models/).
    # Extend this list as more <key>_model.pkl are trained (train_flags_models.py).
    ML_FLAG_KEYS = ['herg', 'dili', 'ames', 'bbb',
                    'cyp3a4', 'cyp2d6', 'cyp2c9', 'bioavail']

    # ADME regression models (value predictions). Clearance (R2~0.18) is
    # intentionally excluded as too weak to be informative.
    ML_ADME_KEYS = ['lipophilicity', 'caco2', 'ppbr']
    
    # Toxicity prediction settings
    TOXICITY_CONFIDENCE_THRESHOLD = 60  # Minimum confidence to trust predictions
    EFFICACY_CONFIDENCE_THRESHOLD = 50
    
    # Request settings
    REQUEST_TIMEOUT = 10
    MAX_RETRIES = 3
    
    # Cache settings
    CACHE_DIR = './cache'
    CACHE_EXPIRY_HOURS = 24
    
    # ML settings
    ML_MODEL_PATH = './ml_models'
    ML_TRAINING_DATA_PATH = './training_data'
    
    # Rate limiting
    RATE_LIMIT_ENABLED = True
    RATE_LIMIT_DEFAULT = "200 per day"
    RATE_LIMIT_PREDICT = "10 per minute"

# ============================================================================
# LOGGING SETUP
# ============================================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================================================
# FLASK APP INITIALIZATION
# ============================================================================

app = Flask(__name__)
app.config.from_object(Config)

# Create necessary directories
os.makedirs(Config.CACHE_DIR, exist_ok=True)
os.makedirs(Config.ML_MODEL_PATH, exist_ok=True)
os.makedirs(Config.ML_TRAINING_DATA_PATH, exist_ok=True)
os.makedirs('./templates', exist_ok=True)
os.makedirs('./static/css', exist_ok=True)
os.makedirs('./static/js', exist_ok=True)

# Rate limiter
if Config.RATE_LIMIT_ENABLED:
    limiter = Limiter(
        app=app,
        key_func=get_remote_address,
        default_limits=[Config.RATE_LIMIT_DEFAULT],
        storage_uri="memory://"
    )
else:
    limiter = None

# ============================================================================
# AUTHENTICATION & ACCESS CONTROL
# ============================================================================
# Email+password login, admin-created accounts, every page behind login.
try:
    from auth import init_auth
    init_auth(app)
    AUTH_ENABLED = True
    logger.info("Authentication enabled (login required)")
except Exception as e:
    AUTH_ENABLED = False
    logger.error(f"Authentication setup failed, app runs WITHOUT login: {e}")

# ============================================================================
# DATA VALIDATION SCHEMAS
# ============================================================================

if VALIDATION_AVAILABLE:
    class GroupSchema(Schema):
        class Meta:
            unknown = EXCLUDE
        
        group_name = fields.Str(required=True, validate=validate.Length(min=1, max=200))
        drug_name = fields.Str(required=True, validate=validate.Length(min=1, max=200))
        dose = fields.Float(validate=validate.Range(min=0, max=10000))
        strain = fields.Str()
        # Confidential / unpublished compound: the name stays inside the platform,
        # the investigator supplies the structure, and the class is searched instead.
        confidential = fields.Boolean()
        smiles = fields.Str(validate=validate.Length(max=1000))
        compound_class = fields.Str(validate=validate.Length(max=200))
        # Drive the sample size from the real endpoint rather than a convention.
        primary_endpoint = fields.Str(validate=validate.Length(max=200))
        detectable_difference_pct = fields.Str(validate=validate.Length(max=20))

        sex = fields.Str(validate=validate.OneOf(['Male', 'Female', 'Mixed']))
        target_organ = fields.Str(validate=validate.Length(max=200))
        num_mice = fields.Integer(required=True, validate=validate.Range(min=1, max=100))
        # Weight range covers both mice (~10-60 g) and rats (~150-600 g).
        weight = fields.Float(validate=validate.Range(min=10, max=600))
        age = fields.Float(validate=validate.Range(min=2, max=104))
        route = fields.Str(validate=validate.OneOf(['Oral', 'IP', 'IV', 'SC', 'IM', 'ICV']))
        diet_type = fields.Str(validate=validate.OneOf(['Standard', 'High fat', 'Fast']))
        start_date = fields.Str()
        sample_types = fields.List(fields.Str())
        blood_quantity = fields.Str()
        # Newer form fields (kept so they pass validation instead of being dropped)
        species = fields.Str(validate=validate.OneOf(['Mouse', 'Rat']))
        experiment_type = fields.Str(validate=validate.Length(max=100))
        toxicity_endpoints = fields.List(fields.Str())
    
    class StudySchema(Schema):
        class Meta:
            unknown = EXCLUDE
        
        study_title = fields.Str(validate=validate.Length(max=500))
        pi_name = fields.Str(validate=validate.Length(max=200))
        groups = fields.List(fields.Nested(GroupSchema), required=True, validate=validate.Length(min=1, max=20))

# ============================================================================
# BLOOD QUANTITY CALCULATOR
# ============================================================================

class BloodQuantityCalculator:
    """Calculate blood quantities needed for various assays."""
    
    # Mouse blood volume: ~70-80 mL/kg (average 75 mL/kg)
    # Safe single collection: 10% of total blood volume
    # Safe repeated collection: 7.5% every 2 weeks
    
    ASSAY_REQUIREMENTS = {
        # Format: (min_volume_ul, optimal_volume_ul, sample_type)
        'Complete Blood Count (CBC)': (50, 100, 'whole blood'),
        'Blood Chemistry Panel': (100, 200, 'serum'),
        'Lipid Profile': (50, 100, 'serum'),
        'Glucose': (5, 10, 'whole blood'),
        'Cytokine Analysis (ELISA)': (50, 100, 'plasma'),
        'Western Blot': (100, 200, 'plasma'),
        'PCR/Gene Expression': (100, 200, 'whole blood'),
        'Flow Cytometry': (50, 100, 'whole blood'),
        'Pharmacokinetics': (25, 50, 'plasma'),
        'Antibody Titers': (50, 100, 'serum'),
        'Coagulation Studies': (200, 300, 'plasma'),
        'Metabolomics': (100, 200, 'plasma'),
        'Proteomics': (200, 300, 'plasma')
    }
    
    @staticmethod
    def calculate_total_blood_volume(weight_g):
        """Calculate total blood volume in mL."""
        weight_kg = weight_g / 1000.0
        return weight_kg * 75  # 75 mL/kg
    
    @staticmethod
    def calculate_safe_collection(weight_g, timepoints=1):
        """Calculate safe collection volume."""
        total_volume = BloodQuantityCalculator.calculate_total_blood_volume(weight_g)
        
        if timepoints == 1:
            # Single collection: 10% of total blood
            safe_volume = total_volume * 0.10
        else:
            # Multiple collections: 7.5% per collection (assuming 2-week intervals)
            safe_volume = total_volume * 0.075
        
        return safe_volume * 1000  # Convert to microliters
    
    @staticmethod
    def calculate_blood_needed(sample_types, weight_g=25, timepoints=1, num_replicates=1):
        """
        Calculate total blood needed for selected sample types.
        
        Args:
            sample_types: List of sample type strings
            weight_g: Mouse weight in grams
            timepoints: Number of blood collection timepoints
            num_replicates: Number of technical replicates per assay
        
        Returns:
            Dictionary with blood requirements and safety assessment
        """
        total_needed = 0
        breakdown = []
        
        # Identify blood-related samples
        blood_samples = [s for s in sample_types if any(keyword in s.lower()
                        for keyword in ['blood', 'plasma', 'serum'])]
        
        if not blood_samples:
            return {
                'needed': False,
                'total_volume_ul': 0,
                'safe_volume_ul': 0,
                'breakdown': [],
                'safety_assessment': 'No blood collection needed'
            }
        
        # One representative assay per blood sample type (a researcher who ticks
        # "serum" usually runs a single chemistry panel — not three assays). This
        # avoids over-estimating the volume and raising a false warning.
        rep_assay = {
            'whole blood': 'Complete Blood Count (CBC)',
            'plasma': 'Pharmacokinetics',
            'serum': 'Blood Chemistry Panel',
        }
        assays_needed = []
        for sample in blood_samples:
            sample_lower = sample.lower()
            for kw, assay in rep_assay.items():
                if kw in sample_lower and assay not in assays_needed:
                    assays_needed.append(assay)
        if not assays_needed:                       # generic "blood" with no subtype
            assays_needed = ['Blood Chemistry Panel']
        
        # Calculate total volume needed
        for assay in assays_needed:
            if assay in BloodQuantityCalculator.ASSAY_REQUIREMENTS:
                min_vol, opt_vol, sample_type = BloodQuantityCalculator.ASSAY_REQUIREMENTS[assay]
                volume_per_replicate = opt_vol
                total_for_assay = volume_per_replicate * num_replicates * timepoints
                
                total_needed += total_for_assay
                breakdown.append({
                    'assay': assay,
                    'volume_per_sample_ul': opt_vol,
                    'replicates': num_replicates,
                    'timepoints': timepoints,
                    'total_ul': total_for_assay,
                    'sample_type': sample_type
                })
        
        # Add 20% overage for pipetting loss
        total_needed_with_overage = total_needed * 1.2
        
        # Calculate safe collection volume
        safe_volume = BloodQuantityCalculator.calculate_safe_collection(weight_g, timepoints)
        total_blood_volume = BloodQuantityCalculator.calculate_total_blood_volume(weight_g)
        # At the study endpoint blood is collected terminally (under anaesthesia,
        # e.g. cardiac puncture), where ~40% of total blood volume is routinely
        # obtainable — far more than the 10% survival-bleed limit. Judge against
        # both so a normal terminal sample isn't flagged as unsafe.
        terminal_safe = total_blood_volume * 1000 * 0.40   # µL

        if total_needed_with_overage <= safe_volume:
            safety = "✓ SAFE - Within the survival (non-terminal) single-collection limit"
            safety_color = "green"
        elif total_needed_with_overage <= terminal_safe:
            safety = "✓ SAFE at terminal collection - Obtainable under terminal anaesthesia (e.g. cardiac puncture) at the study endpoint"
            safety_color = "green"
        elif total_needed_with_overage <= terminal_safe * 1.15:
            safety = "⚠ CAUTION - Near the maximum obtainable volume; collect terminally, reduce assays, or pool across animals"
            safety_color = "orange"
        else:
            safety = "✗ UNSAFE - Exceeds the volume obtainable from one animal. Reduce assays, pool samples, or add animals/timepoints"
            safety_color = "red"
        
        return {
            'needed': True,
            'total_volume_ul': round(total_needed_with_overage, 1),
            'total_volume_ml': round(total_needed_with_overage / 1000, 2),
            'safe_volume_ul': round(safe_volume, 1),
            'safe_volume_ml': round(safe_volume / 1000, 2),
            'total_blood_volume_ml': round(total_blood_volume, 2),
            'percentage_of_total': round((total_needed_with_overage / (total_blood_volume * 1000)) * 100, 1),
            'breakdown': breakdown,
            'safety_assessment': safety,
            'safety_color': safety_color,
            'recommendations': [
                f"Total blood volume in mouse: ~{round(total_blood_volume, 1)} mL",
                f"Safe single collection: {round(safe_volume/1000, 2)} mL ({round(safe_volume, 0)} µL)",
                "Allow 2-week recovery between collections for repeated sampling",
                "Monitor animals for signs of anemia (lethargy, pale feet/ears)",
                "Consider terminal collection if volume exceeds safe limits"
            ]
        }

# ============================================================================
# TOXICITY PREDICTION
# ============================================================================

class ToxicityPredictor:
    """
    Predict drug toxicity using molecular descriptors and external databases.
    Integrates with existing caching and logging systems.
    """
    
    def __init__(self):
        self.cache = None  # Will be set after api_cache is initialized
        # Load the real trained LD50 model (TDC LD50_Zhu, rat oral acute toxicity)
        self.ld50_model = None
        self.ld50_meta = None
        try:
            if os.path.exists(Config.LD50_MODEL_PATH):
                self.ld50_model = joblib.load(Config.LD50_MODEL_PATH)
                if os.path.exists(Config.LD50_META_PATH):
                    with open(Config.LD50_META_PATH) as f:
                        self.ld50_meta = json.load(f)
                logger.info("Loaded trained LD50 model (real experimental data)")
            else:
                logger.warning("Trained LD50 model not found; run train_ld50_model.py")
        except Exception as e:
            logger.error(f"Failed to load LD50 model: {e}")
        # Load the trained aqueous-solubility model (TDC Solubility_AqSolDB)
        self.solubility_model = None
        self.solubility_meta = None
        try:
            if os.path.exists(Config.SOLUBILITY_MODEL_PATH):
                self.solubility_model = joblib.load(Config.SOLUBILITY_MODEL_PATH)
                if os.path.exists(Config.SOLUBILITY_META_PATH):
                    with open(Config.SOLUBILITY_META_PATH) as f:
                        self.solubility_meta = json.load(f)
                logger.info("Loaded trained solubility model")
            else:
                logger.warning("Trained solubility model not found; run train_solubility_model.py")
        except Exception as e:
            logger.error(f"Failed to load solubility model: {e}")
        # Load the trained half-life model (TDC Half_Life_Obach, human PK)
        self.halflife_model = None
        self.halflife_meta = None
        try:
            if os.path.exists(Config.HALFLIFE_MODEL_PATH):
                self.halflife_model = joblib.load(Config.HALFLIFE_MODEL_PATH)
                if os.path.exists(Config.HALFLIFE_META_PATH):
                    with open(Config.HALFLIFE_META_PATH) as f:
                        self.halflife_meta = json.load(f)
                logger.info("Loaded trained half-life model")
            else:
                logger.warning("Trained half-life model not found; run train_halflife_model.py")
        except Exception as e:
            logger.error(f"Failed to load half-life model: {e}")
        # Batch of binary safety/ADME classifiers (hERG, DILI, Ames, BBB, …)
        self.flag_models = {}   # key -> (model, meta)
        for key in Config.ML_FLAG_KEYS:
            try:
                mpath = os.path.join(Config.ML_MODEL_PATH, f'{key}_model.pkl')
                jpath = os.path.join(Config.ML_MODEL_PATH, f'{key}_meta.json')
                if os.path.exists(mpath):
                    meta = json.load(open(jpath)) if os.path.exists(jpath) else {}
                    self.flag_models[key] = (joblib.load(mpath), meta)
            except Exception as e:
                logger.error(f"Failed to load flag model {key}: {e}")
        if self.flag_models:
            logger.info(f"Loaded {len(self.flag_models)} ML flag models: {list(self.flag_models)}")
        # ADME regression models (Lipophilicity, Caco-2, PPBR, …)
        self.adme_models = {}   # key -> (model, meta)
        for key in Config.ML_ADME_KEYS:
            try:
                mpath = os.path.join(Config.ML_MODEL_PATH, f'{key}_model.pkl')
                jpath = os.path.join(Config.ML_MODEL_PATH, f'{key}_meta.json')
                if os.path.exists(mpath):
                    meta = json.load(open(jpath)) if os.path.exists(jpath) else {}
                    self.adme_models[key] = (joblib.load(mpath), meta)
            except Exception as e:
                logger.error(f"Failed to load ADME model {key}: {e}")
        if self.adme_models:
            logger.info(f"Loaded {len(self.adme_models)} ML ADME models: {list(self.adme_models)}")
        # Applicability-domain references: the training fingerprints of each
        # model, used to withhold predictions for compounds outside the
        # chemical space the model actually learned from.
        self.domain_refs = {}
        try:
            import numpy as _np
            for _f in os.listdir(Config.ML_MODEL_PATH):
                if _f.endswith('_domain.npz'):
                    _k = _f[:-len('_domain.npz')]
                    self.domain_refs[_k] = _np.load(
                        os.path.join(Config.ML_MODEL_PATH, _f))['fps']
        except Exception as e:
            logger.error(f"Failed to load applicability-domain references: {e}")
        if self.domain_refs:
            logger.info(f"Loaded {len(self.domain_refs)} applicability-domain references")
        logger.info("ToxicityPredictor initialized")
    
    def get_chemical_structure(self, drug_name):
        """Fetch SMILES structure from PubChem."""
        cache_key = {'drug': drug_name.lower(), 'type': 'smiles'}
        if self.cache:
            cached = self.cache.get('pubchem_smiles', cache_key)
            if cached:
                return cached
        
        try:
            # Get CID from existing function; propagate its specific error
            # (spelling vs. temporarily-unreachable) instead of masking it.
            cid_data = get_drug_data_from_ncbi(drug_name)
            if not cid_data['success']:
                return {'success': False,
                        'error': cid_data.get('error', 'Drug not found in PubChem')}
            
            cid = cid_data.get('cid')
            if not cid:
                return {'success': False, 'error': 'No CID found'}
            
            # Fetch SMILES. PubChem deprecated 'CanonicalSMILES'; the current
            # fields are 'SMILES' and 'ConnectivitySMILES'. Request all and use
            # whichever is present for backward/forward compatibility.
            url = (f"{Config.PUBCHEM_API_BASE}/compound/cid/{cid}"
                   f"/property/SMILES,ConnectivitySMILES,CanonicalSMILES,IsomericSMILES/JSON")
            response = requests.get(url, timeout=Config.REQUEST_TIMEOUT)
            response.raise_for_status()

            data = response.json()
            props = data['PropertyTable']['Properties'][0]
            smiles = (props.get('SMILES') or props.get('CanonicalSMILES')
                      or props.get('IsomericSMILES') or props.get('ConnectivitySMILES'))
            if not smiles:
                return {'success': False, 'error': 'No SMILES returned by PubChem', 'cid': cid}

            result = {'success': True, 'smiles': smiles, 'cid': cid}
            if self.cache:
                self.cache.set('pubchem_smiles', cache_key, result)
            return result
            
        except Exception as e:
            logger.error(f"Error fetching SMILES for {drug_name}: {e}")
            return {'success': False, 'error': str(e)}
    
    
    
    # ------------------------------------------------------------------
    # LD50 prediction — HYBRID: real experimental data first, trained ML
    # model second, molecular heuristic only as a last resort.
    # Categories (mg/kg): high (<50), moderate (50-500), low (500-2000),
    # very_low (>2000).
    # ------------------------------------------------------------------

    @staticmethod
    def ld50_mgkg_to_category(ld50_mg_kg):
        """Map an LD50 value in mg/kg to a toxicity category + a tight range."""
        if ld50_mg_kg < 50:
            return 'high', (1, 50)
        elif ld50_mg_kg < 500:
            return 'moderate', (50, 500)
        elif ld50_mg_kg < 2000:
            return 'low', (500, 2000)
        else:
            return 'very_low', (2000, 99999)

    def _featurize_smiles(self, smiles, meta=None):
        """Build the feature vector for a trained model. Defaults to the LD50
        meta; the solubility model shares the same descriptor + fingerprint
        layout, so the same featurizer serves both."""
        meta = meta or self.ld50_meta
        if not RDKIT_AVAILABLE or not meta:
            return None
        mol = Chem.MolFromSmiles(smiles)
        if mol is None:
            return None
        desc = {
            'molecular_weight': Descriptors.MolWt(mol),
            'logp': Crippen.MolLogP(mol),
            'h_donors': Lipinski.NumHDonors(mol),
            'h_acceptors': Lipinski.NumHAcceptors(mol),
            'rotatable_bonds': Lipinski.NumRotatableBonds(mol),
            'aromatic_rings': Lipinski.NumAromaticRings(mol),
            'tpsa': Descriptors.TPSA(mol),
        }
        desc_vec = [desc[n] for n in meta['descriptor_names']]
        fp = rdMolDescriptors.GetMorganFingerprintAsBitVect(
            mol, radius=meta['fp_radius'], nBits=meta['fp_bits'])
        return [desc_vec + list(fp)], Descriptors.MolWt(mol)

    def predict_ld50_ml(self, smiles):
        """Predict LD50 (mg/kg) with the trained model. Returns dict or None."""
        if self.ld50_model is None:
            return None
        feats = self._featurize_smiles(smiles)
        if feats is None:
            return None
        X, mw = feats
        try:
            y = float(self.ld50_model.predict(X)[0])  # -log10(mol/kg)
        except Exception as e:
            logger.error(f"LD50 model prediction failed: {e}")
            return None
        dom = self.domain_check(smiles, 'ld50')
        if dom and not dom['in_domain']:
            return None            # no defensible estimate outside the domain
        mol_per_kg = 10 ** (-y)
        ld50_mg_kg = mol_per_kg * mw * 1000
        category, ld50_range = self.ld50_mgkg_to_category(ld50_mg_kg)
        # Report the model's measured held-out performance, not a derived
        # "confidence %" — the test R2 and MAE are the numbers a reviewer can
        # actually check against the published benchmark.
        r2 = (self.ld50_meta or {}).get('test_r2', 0.5)
        mae = (self.ld50_meta or {}).get('test_mae', 0.48)
        return {
            'category': category,
            'ld50_range': ld50_range,
            'ld50_mg_kg': round(ld50_mg_kg, 1),
            'test_r2': round(r2, 2),
            'test_mae_log': round(mae, 2),
            'error_band': f"×/÷ {10 ** mae:.1f}",
            'source': 'ml_model',
            'source_detail': (self.ld50_meta or {}).get('source', 'trained model'),
        }

    # A Morgan/Tanimoto similarity below this to every training compound means
    # the molecule is unlike anything the model learned from, so its output is
    # not meaningful. 0.30 is the conventional cut-off for ECFP-style prints.
    DOMAIN_MIN_SIMILARITY = 0.30

    def domain_check(self, smiles, key):
        """Max Tanimoto similarity of `smiles` to the model's training set.

        Returns {'in_domain': bool, 'similarity': float} or None when the check
        cannot run (missing reference or unparseable structure)."""
        ref = self.domain_refs.get(key)
        if ref is None or not RDKIT_AVAILABLE:
            return None
        try:
            import numpy as np
            mol = Chem.MolFromSmiles(smiles)
            if mol is None:
                return None
            fp = rdMolDescriptors.GetMorganFingerprintAsBitVect(mol, radius=2, nBits=256)
            q = np.packbits(np.array(fp, dtype=np.uint8))
            # Tanimoto on packed bits: |A&B| / (|A| + |B| - |A&B|)
            inter = np.unpackbits(np.bitwise_and(ref, q), axis=1).sum(axis=1)
            a = np.unpackbits(q).sum()
            b = np.unpackbits(ref, axis=1).sum(axis=1)
            union = a + b - inter
            sim = float(np.max(np.where(union > 0, inter / np.maximum(union, 1), 0.0)))
            return {'in_domain': sim >= self.DOMAIN_MIN_SIMILARITY,
                    'similarity': round(sim, 2)}
        except Exception as e:
            logger.warning(f"Domain check failed for {key}: {e}")
            return None

    @staticmethod
    def logs_to_solubility(logs):
        """Map predicted logS (log10 mol/L) to a class + a vehicle/formulation hint."""
        if logs >= -2:
            return 'high', 'Aqueous vehicle (saline / PBS) is typically adequate.'
        elif logs >= -4:
            return 'moderate', 'May need a co-solvent (e.g. 5-10% DMSO or cyclodextrin) or gentle warming.'
        else:
            return 'low', ('Poorly water-soluble — consider DMSO / cyclodextrin / a cosolvent '
                           'system, a nanosuspension, or a salt form.')

    def predict_solubility_ml(self, smiles, feats=None):
        """Predict aqueous solubility (logS) with the trained model. Returns a
        dict {logS, mg_per_mL, category, vehicle_hint, confidence, source} or None.
        `feats` (from a shared featurization) is reused if provided."""
        if self.solubility_model is None:
            return None
        feats = feats if feats is not None else self._featurize_smiles(smiles, self.solubility_meta)
        if feats is None:
            return None
        X, mw = feats
        try:
            logs = float(self.solubility_model.predict(X)[0])   # log10(mol/L)
        except Exception as e:
            logger.error(f"Solubility model prediction failed: {e}")
            return None
        dom = self.domain_check(smiles, 'solubility')
        if dom and not dom['in_domain']:
            return {'out_of_domain': True, 'similarity': dom['similarity'],
                    'source': 'ml_model'}
        category, vehicle_hint = self.logs_to_solubility(logs)
        mg_per_mL = (10 ** logs) * mw          # mol/L * g/mol = g/L = mg/mL
        r2 = (self.solubility_meta or {}).get('test_r2', 0.78)
        mae = (self.solubility_meta or {}).get('test_mae', 0.76)
        return {
            'logS': round(logs, 2),
            'mg_per_mL': round_sig(mg_per_mL, 2),
            'category': category,              # high / moderate / low
            'vehicle_hint': vehicle_hint,
            'test_r2': round(r2, 2),
            'test_mae_log': round(mae, 2),
            'error_band': f"×/÷ {10 ** mae:.1f}",
            'source': 'ml_model',
            'source_detail': (self.solubility_meta or {}).get('source', 'trained model'),
        }

    @staticmethod
    def hours_to_frequency(hours):
        """Map a predicted half-life (hours) to a class + dosing-frequency hint."""
        if hours < 3:
            return 'short', ('Short half-life — frequent dosing (e.g. twice/three-times '
                             'daily) or continuous delivery may be needed.')
        elif hours < 12:
            return 'intermediate', 'Once-to-twice daily dosing is often appropriate.'
        else:
            return 'long', 'Long half-life — once daily (or less frequent) dosing may suffice.'

    def predict_halflife_ml(self, smiles, feats=None):
        """Predict plasma half-life with the trained model. NOTE: trained on HUMAN
        data (Obach) and is a modest model — used as a *relative* indicator for
        dosing frequency, not an absolute rodent value. Returns dict or None.
        `feats` (from a shared featurization) is reused if provided."""
        if self.halflife_model is None:
            return None
        feats = feats if feats is not None else self._featurize_smiles(smiles, self.halflife_meta)
        if feats is None:
            return None
        X, _mw = feats
        try:
            y = float(self.halflife_model.predict(X)[0])   # log10(hours)
        except Exception as e:
            logger.error(f"Half-life model prediction failed: {e}")
            return None
        hours = 10 ** y
        dom = self.domain_check(smiles, 'halflife')
        if dom and not dom['in_domain']:
            return {'out_of_domain': True, 'similarity': dom['similarity'],
                    'source': 'ml_model'}
        category, frequency_hint = self.hours_to_frequency(hours)
        r2 = (self.halflife_meta or {}).get('test_r2', 0.27)
        mae = (self.halflife_meta or {}).get('test_mae', 0.40)
        # This model explains only ~27% of the variance (test MAE 0.40 log units
        # ~ a factor of 2.5), so the hour value is NOT reported — only the
        # ordinal band, which is what the dosing-frequency hint actually needs.
        return {
            'hours': None,
            'value_withheld': True,
            'withheld_reason': (
                f"Model too weak to quote a number (test R² {r2:.2f}, "
                f"MAE {mae:.2f} log units ≈ a factor of {10 ** mae:.1f}). "
                "Only the ordinal band is reported."),
            'category': category,          # short / intermediate / long
            'frequency_hint': frequency_hint,
            'test_r2': round(r2, 2),
            'species': (self.halflife_meta or {}).get('species', 'human'),
            'source': 'ml_model',
            'source_detail': (self.halflife_meta or {}).get('source', 'trained model'),
        }

    # Which flag keys are safety RISKS (high probability = concern) vs neutral
    # structural properties (e.g. BBB penetration / bioavailability are
    # informative, not risks).
    _RISK_FLAG_KEYS = {'herg', 'dili', 'ames', 'cyp3a4', 'cyp2d6', 'cyp2c9'}

    def predict_flags_ml(self, smiles, feats=None):
        """Run every loaded binary safety/ADME classifier on a SMILES.
        All classifiers share one feature layout, so the molecule is featurized
        ONCE (or `feats` is reused) — keeps this O(1) in RDKit work as the model
        registry grows. Returns a list of flag dicts or []."""
        if not self.flag_models:
            return []
        shared = feats if feats is not None else self._featurize_smiles(smiles)
        if shared is None:
            return []
        X, _mw = shared
        out = []
        for key, (model, meta) in self.flag_models.items():
            try:
                p = float(model.predict_proba(X)[0][1])
            except Exception as e:
                logger.error(f"Flag model {key} prediction failed: {e}")
                continue
            dom = self.domain_check(smiles, key)
            if dom and not dom['in_domain']:
                # Outside the model's chemical space — report that, not a number.
                out.append({
                    'key': key, 'label': meta.get('label', key),
                    'out_of_domain': True, 'similarity': dom['similarity'],
                    'is_risk': key in self._RISK_FLAG_KEYS,
                    'auc': round(meta.get('test_auc', 0.75), 2),
                    'source': meta.get('source', ''),
                })
                continue
            auc = round(meta.get('test_auc', 0.75), 2)
            out.append({
                'key': key,
                'label': meta.get('label', key),
                'probability': round(p, 2),
                'domain_similarity': (dom or {}).get('similarity'),
                'flag': 'high' if p >= 0.5 else 'low',
                'is_risk': key in self._RISK_FLAG_KEYS,
                'meaning': meta.get('positive_meaning', ''),
                'auc': auc,
                'reliability': self.auc_reliability(auc),
                'source': meta.get('source', ''),
            })
        return out

    @staticmethod
    def auc_reliability(auc):
        """Tier a classifier by its held-out AUC so weak models are not shown
        with the same weight as strong ones."""
        if auc >= 0.85:
            return 'good'
        if auc >= 0.75:
            return 'moderate'
        return 'weak'

    @staticmethod
    def r2_reliability(r2):
        """Tier a regression model by its held-out R²."""
        if r2 >= 0.70:
            return 'good'
        if r2 >= 0.50:
            return 'moderate'
        return 'weak'

    def predict_adme_ml(self, smiles, feats=None):
        """Run every loaded ADME regression model on a SMILES (featurized once /
        `feats` reused). Returns a list of {label, value, unit, r2} or []."""
        if not self.adme_models:
            return []
        shared = feats if feats is not None else self._featurize_smiles(smiles)
        if shared is None:
            return []
        X, _mw = shared
        out = []
        for key, (model, meta) in self.adme_models.items():
            try:
                p = float(model.predict(X)[0])
            except Exception as e:
                logger.error(f"ADME model {key} prediction failed: {e}")
                continue
            dom = self.domain_check(smiles, key)
            if dom and not dom['in_domain']:
                out.append({'key': key, 'label': meta.get('label', key),
                            'out_of_domain': True, 'similarity': dom['similarity']})
                continue
            val = 10 ** p if meta.get('log') else p
            r2 = round(meta.get('test_r2', 0.5), 2)
            mae = meta.get('test_mae')
            out.append({
                'key': key,
                'label': meta.get('label', key),
                'value': round_sig(val, 2),
                'unit': meta.get('unit', ''),
                'r2': r2,
                'reliability': self.r2_reliability(r2),
                'test_mae': round(mae, 2) if mae is not None else None,
            })
        return out

    def _parse_dose_mgkg(self, dose_str):
        """Parse a ChemIDplus dose string like '250 mg/kg' -> 250.0 (mg/kg)."""
        if not dose_str:
            return None
        m = re.match(r'([\d.]+)\s*(mg|gm|g|ug|mcg|ng)/kg', dose_str.strip(), re.I)
        if not m:
            return None
        val = float(m.group(1))
        unit = m.group(2).lower()
        factor = {'mg': 1, 'gm': 1000, 'g': 1000, 'ug': 1e-3, 'mcg': 1e-3, 'ng': 1e-6}[unit]
        return val * factor

    def get_experimental_ld50(self, drug_name, cid=None, route='oral', species='Mouse'):
        """
        Fetch REAL experimental LD50 values from PubChem's ChemIDplus records.
        Prefers the requested species (Mouse or Rat), then the other rodent,
        and matching administration route.
        Returns dict with 'found' and (if found) the best LD50 in mg/kg.
        """
        preferred = (species or 'Mouse').strip().lower()
        if preferred not in ('mouse', 'rat'):
            preferred = 'mouse'
        cache_key = {'drug': drug_name.lower(), 'route': route.lower(),
                     'species': preferred, 'type': 'exp_ld50'}
        if self.cache:
            cached = self.cache.get('experimental_ld50', cache_key)
            if cached:
                return cached

        result = {'found': False}
        try:
            if cid is None:
                cid_data = get_drug_data_from_ncbi(drug_name)
                cid = cid_data.get('cid') if cid_data.get('success') else None
            if not cid:
                return result

            # Step 1: find the ChemIDplus SID from the Acute Effects section
            try:
                resp = requests.get(
                    f"{Config.PUBCHEM_PUGVIEW_BASE}/{cid}/JSON",
                    params={'heading': 'Acute Effects'},
                    timeout=Config.REQUEST_TIMEOUT)
                resp.raise_for_status()
                record = resp.json().get('Record', {})
            except requests.exceptions.HTTPError:
                # 404 = no acute-effects data for this compound -> fall back to model
                if self.cache:
                    self.cache.set('experimental_ld50', cache_key, result)
                return result

            def find_sid(sec):
                for s in sec.get('Section', []):
                    for info in s.get('Information', []):
                        ext = info.get('Value', {}).get('ExternalTableName', '')
                        m = re.search(r'query=(\d+)', ext)
                        if m:
                            return m.group(1)
                    found = find_sid(s)
                    if found:
                        return found
                return None

            sid = find_sid(record)
            if not sid:
                if self.cache:
                    self.cache.set('experimental_ld50', cache_key, result)
                return result

            # Step 2: pull the ChemIDplus toxicity rows via the SDQ endpoint
            query = ('{"select":"*","collection":"chemidplus","where":{"ands":'
                     '[{"sid":"%s"}]},"order":["relevancescore,desc"],'
                     '"start":1,"limit":80}') % sid
            resp = requests.get(
                Config.PUBCHEM_SDQ,
                params={'infmt': 'json', 'outfmt': 'json', 'query': query},
                timeout=Config.REQUEST_TIMEOUT)
            resp.raise_for_status()
            data = resp.json()
            rows = data if isinstance(data, list) else \
                data.get('SDQOutputSet', [{}])[0].get('rows', [])

            ld50_entries = []
            for row in rows:
                if row.get('testtype') != 'LD50':
                    continue
                mg = self._parse_dose_mgkg(row.get('dose', ''))
                if mg is None or mg <= 0:
                    continue
                ld50_entries.append({
                    'organism': row.get('organism', ''),
                    'route': row.get('route', ''),
                    'ld50_mg_kg': round(mg, 2),
                })

            if not ld50_entries:
                if self.cache:
                    self.cache.set('experimental_ld50', cache_key, result)
                return result

            # Step 3: rank — mouse > rat, and matching route preferred.
            # Map short route codes to the terms ChemIDplus actually uses.
            route_synonyms = {
                'oral': ['oral'],
                'iv': ['intravenous'], 'intravenous': ['intravenous'],
                'ip': ['intraperitoneal'], 'intraperitoneal': ['intraperitoneal'],
                'sc': ['subcutaneous'], 'subcutaneous': ['subcutaneous'],
                'im': ['intramuscular'], 'intramuscular': ['intramuscular'],
                'icv': ['intracerebral', 'intracranial'],
            }
            wanted_routes = route_synonyms.get(route.lower(), [route.lower()])

            other_rodent = 'rat' if preferred == 'mouse' else 'mouse'

            def rank(e):
                s = 0
                org = e['organism'].lower()
                if org == preferred:
                    s += 4            # the species the user selected
                elif org == other_rodent:
                    s += 3            # the other rodent
                er = (e['route'] or '').lower()
                if er and any(w in er for w in wanted_routes):
                    s += 2
                return -s
            ld50_entries.sort(key=rank)
            best = ld50_entries[0]
            category, ld50_range = self.ld50_mgkg_to_category(best['ld50_mg_kg'])

            result = {
                'found': True,
                'category': category,
                'ld50_range': ld50_range,
                'ld50_mg_kg': best['ld50_mg_kg'],
                'confidence': 95,  # real experimental value
                'source': 'experimental',
                'source_detail': f"PubChem/ChemIDplus ({best['organism']}, {best['route']})",
                'best': best,
                'all_values': ld50_entries[:12],
            }
        except Exception as e:
            logger.error(f"Experimental LD50 lookup failed for {drug_name}: {e}")
            return {'found': False}

        if self.cache:
            self.cache.set('experimental_ld50', cache_key, result)
        return result


    
    def get_safe_dose_recommendations(self, ld50_category, route='oral'):
        """Calculate safe starting doses based on LD50 category."""
        dose_guidelines = {
            'very_low': {
                'starting_dose': '10-50 mg/kg',
                'max_dose': '500 mg/kg',
                'escalation': 'Can escalate rapidly (2x per week)',
                'monitoring': 'Standard monitoring sufficient'
            },
            'low': {
                'starting_dose': '5-20 mg/kg',
                'max_dose': '100 mg/kg',
                'escalation': 'Moderate escalation (1.5x per week)',
                'monitoring': 'Weekly health checks recommended'
            },
            'moderate': {
                'starting_dose': '1-5 mg/kg',
                'max_dose': '20 mg/kg',
                'escalation': 'Slow escalation (1.2x per week)',
                'monitoring': 'Daily health checks required'
            },
            'high': {
                'starting_dose': '0.1-1 mg/kg',
                'max_dose': '5 mg/kg',
                'escalation': 'Very slow escalation (1.1x per week)',
                'monitoring': 'Continuous monitoring required + veterinary oversight'
            }
        }
        
        return dose_guidelines.get(ld50_category, dose_guidelines['moderate'])
    
    
    

    def generate_welfare_recommendations(self, risk_score, organ_risks=None,
                                         pct_ld50=0, target_organ=None):
        """
        Derive HUMANE ENDPOINTS (IACUC Part 8) and PAIN/DISTRESS MONITORING
        (Part 11) from the toxicity assessment. These map directly onto the
        animal-ethics protocol so a researcher can reuse them.

        DECISION-SUPPORT ONLY — must be reviewed by the researcher and the
        attending veterinarian. Does NOT cover the euthanasia METHOD (that is
        institutionally standardised and out of scope).
        """
        organ_risks = dict(organ_risks or {})
        # If only a single target organ is known (comprehensive endpoint),
        # infer its risk from the overall score.
        if not organ_risks and target_organ and target_organ.lower() not in ('general', '', 'none'):
            level = 'high' if risk_score >= 60 else 'moderate' if risk_score >= 40 else 'low'
            organ_risks = {target_organ.lower(): level}

        # --- Humane endpoints (Part 8) ---
        humane_endpoints = [
            "Body weight loss ≥ 20% from baseline (or ≥ 15% with other clinical signs)",
            "Body condition score ≤ 2/5 (emaciation)",
            "Prolonged anorexia or inability to reach food/water (>24 h)",
            "Persistent hypothermia, hunched posture, or lack of grooming",
        ]
        organ_signs = {
            'liver': "hepatotoxicity signs — jaundice, marked lethargy, abnormal bleeding",
            'kidney': "nephrotoxicity signs — reduced urination, oedema, dehydration",
            'brain': "neurological signs — seizures, tremors, paralysis, circling, ataxia",
            'heart': "cardiovascular signs — cyanosis, laboured breathing, collapse",
            'lung': "respiratory distress — dyspnoea, gasping, cyanosis",
        }
        for organ, risk in organ_risks.items():
            key = organ.lower()
            if risk in ('high', 'moderate') and key in organ_signs:
                humane_endpoints.append(
                    f"[{key.upper()} risk: {risk}] Monitor for {organ_signs[key]}")

        if pct_ld50 and pct_ld50 >= 25:
            humane_endpoints.append(
                f"Dose is ~{round(pct_ld50, 1)}% of the estimated LD50 — watch closely for "
                f"acute toxicity; euthanise on moderate–severe distress that does not resolve.")

        # --- Pain / distress monitoring (Part 11) ---
        if risk_score >= 60:
            level = 'high'
            frequency = "At least twice daily (≈12 h apart); continuous around peak drug effect."
        elif risk_score >= 40:
            level = 'moderate'
            frequency = "Once daily, with additional checks after each dosing event."
        else:
            level = 'low'
            frequency = "Every 2–3 days; daily around dosing events."

        if risk_score >= 40:
            analgesia = ("Provide analgesia/anti-inflammatory as advised by the veterinarian; "
                         "do not withhold without scientific justification (Part 11).")
        else:
            analgesia = ("Analgesia likely not required for routine handling/injection; "
                         "reassess promptly if any distress is observed.")

        return {
            'disclaimer': ("Decision-support suggestions derived from predicted toxicity. "
                           "Must be reviewed and approved by the researcher and attending "
                           "veterinarian before use in a protocol."),
            'humane_endpoints': humane_endpoints,
            'monitoring': {
                'level': level,
                'frequency': frequency,
                'analgesia_guidance': analgesia,
                'body_condition_scoring': ("Record body weight and body condition score (1–5) "
                                           "at each observation."),
            },
            'response_when_endpoint_reached': (
                "Remove the animal from study and apply the approved euthanasia method "
                "immediately; record observations and notify the veterinarian."),
        }

    


# ============================================================================
# EFFECTIVENESS PREDICTION
# ============================================================================



# ============================================================================
# DECISION SUPPORT SYSTEM
# ============================================================================









# ============================================================================
# CACHING SYSTEM
# ============================================================================

class APICache:
    """Simple file-based cache for API responses."""
    
    def __init__(self, cache_dir=Config.CACHE_DIR, ttl_hours=Config.CACHE_EXPIRY_HOURS):
        self.cache_dir = cache_dir
        self.ttl = timedelta(hours=ttl_hours)
        os.makedirs(cache_dir, exist_ok=True)
    
    def _get_cache_key(self, prefix, params):
        """Generate cache key from parameters."""
        key_string = f"{prefix}_{str(sorted(params.items()))}"
        return hashlib.md5(key_string.encode()).hexdigest()
    
    def get(self, prefix, params):
        """Get cached data if valid."""
        try:
            cache_key = self._get_cache_key(prefix, params)
            cache_file = os.path.join(self.cache_dir, f"{cache_key}.pkl")
            
            if os.path.exists(cache_file):
                with open(cache_file, 'rb') as f:
                    cached = pickle.load(f)
                    if datetime.now() - cached['timestamp'] < self.ttl:
                        logger.debug(f"Cache hit for {prefix}")
                        return cached['data']
                    else:
                        os.remove(cache_file)
        except Exception as e:
            logger.warning(f"Cache read error: {e}")
        return None
    
    def set(self, prefix, params, data):
        """Store data in cache."""
        try:
            cache_key = self._get_cache_key(prefix, params)
            cache_file = os.path.join(self.cache_dir, f"{cache_key}.pkl")
            
            with open(cache_file, 'wb') as f:
                pickle.dump({
                    'timestamp': datetime.now(),
                    'data': data
                }, f)
            logger.debug(f"Cached data for {prefix}")
        except Exception as e:
            logger.warning(f"Cache write error: {e}")

api_cache = APICache()

# Initialize prediction engines
toxicity_predictor = ToxicityPredictor()
toxicity_predictor.cache = api_cache  # Connect to cache system

# ============================================================================
# RATE LIMITING FOR EXTERNAL APIs
# ============================================================================

class RateLimiter:
    """Simple rate limiter for API calls."""
    
    def __init__(self, calls_per_second=3):
        self.calls_per_second = calls_per_second
        self.last_call = 0
    
    def wait(self):
        """Wait if necessary to respect rate limit."""
        elapsed = time.time() - self.last_call
        wait_time = (1.0 / self.calls_per_second) - elapsed
        if wait_time > 0:
            time.sleep(wait_time)
        self.last_call = time.time()

ncbi_rate_limiter = RateLimiter(calls_per_second=3)
semantic_rate_limiter = RateLimiter(calls_per_second=2)
openalex_rate_limiter = RateLimiter(calls_per_second=2)

# ============================================================================
# MACHINE LEARNING MODEL
# ============================================================================


# Initialize ML model
ml_model = None  # retired: sample size now comes from power analysis only

# ============================================================================
# ENHANCED API FUNCTIONS - MULTIPLE SOURCES
# ============================================================================

def is_confidential(group):
    """True when the investigator marked the compound as unpublished/confidential."""
    return bool((group or {}).get('confidential'))


def external_term(group):
    """The compound term used in EXTERNAL queries.

    For a confidential compound the name never leaves the platform: its class
    (e.g. 'anthracycline') is searched instead, so the literature is retrieved
    from the study context rather than from an identifying name.
    """
    group = group or {}
    if is_confidential(group):
        return (group.get('compound_class') or '').strip()
    return (group.get('drug_name') or '').strip()


def search_semantic_scholar(group, max_results=5):
    """Search Semantic Scholar API for papers."""
    cache_key = {
        'drug': external_term(group),
        'strain': group.get('strain', ''),
        'target': group.get('target_organ', ''),
        'max': max_results
    }
    
    cached = api_cache.get('semantic_scholar', cache_key)
    if cached:
        return cached
    
    query_parts = [
        external_term(group),
        group.get('strain', ''),
        group.get('target_organ', ''),
        'mouse model'
    ]
    query = ' '.join([p for p in query_parts if p]).strip()
    
    if not query:
        return []
    
    try:
        semantic_rate_limiter.wait()
        
        url = f"{Config.SEMANTIC_SCHOLAR_BASE}/paper/search"
        params = {
            'query': query,
            'limit': max_results,
            'fields': 'title,authors,year,abstract,citationCount,venue,publicationDate,externalIds'
        }
        
        response = requests.get(url, params=params, timeout=Config.REQUEST_TIMEOUT)
        response.raise_for_status()
        
        data = response.json()
        papers = []
        
        for paper in data.get('data', []):
            external_ids = paper.get('externalIds', {})
            doi = external_ids.get('DOI', '')
            pmid = external_ids.get('PubMed', '')
            
            papers.append({
                'title': paper.get('title', ''),
                'authors': ', '.join([a.get('name', '') for a in paper.get('authors', [])[:3]]),
                'year': paper.get('year'),
                'venue': paper.get('venue', ''),
                'citations': paper.get('citationCount', 0),
                'doi': doi,
                'pmid': pmid,
                'url': f"https://www.semanticscholar.org/paper/{paper.get('paperId', '')}",
                'source': 'Semantic Scholar'
            })
        
        api_cache.set('semantic_scholar', cache_key, papers)
        return papers
        
    except Exception as e:
        logger.error(f"Semantic Scholar search failed: {e}")
        return []

def search_openalex(group, max_results=5):
    """Search OpenAlex API for papers."""
    cache_key = {
        'drug': external_term(group),
        'strain': group.get('strain', ''),
        'target': group.get('target_organ', ''),
        'max': max_results
    }
    
    cached = api_cache.get('openalex', cache_key)
    if cached:
        return cached
    
    query_parts = [
        external_term(group),
        group.get('strain', ''),
        group.get('target_organ', ''),
        'mouse'
    ]
    query = ' '.join([p for p in query_parts if p]).strip()
    
    if not query:
        return []
    
    try:
        openalex_rate_limiter.wait()
        
        url = f"{Config.OPENALEX_BASE}/works"
        params = {
            'search': query,
            'filter': 'type:journal-article',
            'per-page': max_results,
            'mailto': 'research@example.com'  # Polite pool access
        }
        
        response = requests.get(url, params=params, timeout=Config.REQUEST_TIMEOUT)
        response.raise_for_status()
        
        data = response.json()
        papers = []
        
        for work in data.get('results', []):
            papers.append({
                'title': work.get('title', ''),
                'authors': ', '.join([a.get('author', {}).get('display_name', '')
                                     for a in work.get('authorships', [])[:3]]),
                'year': work.get('publication_year'),
                'venue': work.get('primary_location', {}).get('source', {}).get('display_name', ''),
                'citations': work.get('cited_by_count', 0),
                'doi': work.get('doi', '').replace('https://doi.org/', ''),
                'url': work.get('doi', work.get('id', '')),
                'source': 'OpenAlex'
            })
        
        api_cache.set('openalex', cache_key, papers)
        return papers
        
    except Exception as e:
        logger.error(f"OpenAlex search failed: {e}")
        return []

def search_crossref(group, max_results=5):
    """Search CrossRef API for papers."""
    cache_key = {
        'drug': external_term(group),
        'strain': group.get('strain', ''),
        'target': group.get('target_organ', ''),
        'max': max_results
    }
    
    cached = api_cache.get('crossref', cache_key)
    if cached:
        return cached
    
    query_parts = [
        external_term(group),
        group.get('strain', ''),
        group.get('target_organ', ''),
        'mouse study'
    ]
    query = ' '.join([p for p in query_parts if p]).strip()
    
    if not query:
        return []
    
    try:
        url = Config.CROSSREF_BASE
        params = {
            'query': query,
            'rows': max_results,
            'filter': 'type:journal-article',
            'mailto': 'research@example.com'
        }
        
        response = requests.get(url, params=params, timeout=Config.REQUEST_TIMEOUT)
        response.raise_for_status()
        
        data = response.json()
        papers = []
        
        for item in data.get('message', {}).get('items', []):
            # Extract authors
            authors = []
            for author in item.get('author', [])[:3]:
                given = author.get('given', '')
                family = author.get('family', '')
                authors.append(f"{given} {family}".strip())
            
            # Extract publication date
            pub_date = item.get('published-print', item.get('published-online', {}))
            year = pub_date.get('date-parts', [[None]])[0][0] if pub_date else None
            
            papers.append({
                'title': item.get('title', [''])[0],
                'authors': ', '.join(authors),
                'year': year,
                'venue': item.get('container-title', [''])[0],
                'doi': item.get('DOI', ''),
                'url': f"https://doi.org/{item.get('DOI', '')}",
                'source': 'CrossRef'
            })
        
        api_cache.set('crossref', cache_key, papers)
        return papers
        
    except Exception as e:
        logger.error(f"CrossRef search failed: {e}")
        return []

# [Previous NCBI, Europe PMC, and IMPC functions remain the same]

def get_drug_data_from_ncbi(drug_name):
    """Get PubChem CID and build compound URL."""
    if not drug_name:
        return {"success": False, "error": "Drug name is required."}

    # Normalise: trim surrounding whitespace so a stray trailing space does not
    # turn a valid name into a 404 (e.g. "ibuprofen " -> ".../ibuprofen%20/...").
    drug_name = drug_name.strip()
    if not drug_name:
        return {"success": False, "error": "Drug name is required."}

    if drug_name.lower() in ['saline', 'control', 'pbs', 'vehicle']:
        return {"success": True, "ncbi_link": "#", "is_control": True}

    cache_key = {'drug_name': drug_name.lower()}
    cached = api_cache.get('pubchem', cache_key)
    if cached:
        return cached

    try:
        # URL-encode the name so spaces/special characters don't break the URL.
        encoded = requests.utils.quote(drug_name, safe='')
        url = f"{Config.PUBCHEM_API_BASE}/compound/name/{encoded}/cids/JSON"
        response = requests.get(url, timeout=Config.REQUEST_TIMEOUT)

        # 404 = the drug name is genuinely not in PubChem (often a typo),
        # NOT a connectivity problem. Give an accurate, actionable message.
        if response.status_code == 404:
            result = {"success": False,
                      "error": f"Drug '{drug_name}' not found in PubChem — please check the spelling."}
            api_cache.set('pubchem', cache_key, result)
            return result

        response.raise_for_status()

        cid_list = response.json().get('IdentifierList', {}).get('CID', [])
        if not cid_list:
            result = {"success": False,
                      "error": f"Drug '{drug_name}' not found in PubChem — please check the spelling."}
        else:
            cid = cid_list[0]
            result = {
                "success": True,
                "ncbi_link": f"https://pubchem.ncbi.nlm.nih.gov/compound/{cid}",
                "cid": cid
            }

        api_cache.set('pubchem', cache_key, result)
        return result

    except requests.exceptions.HTTPError as e:
        # Any other HTTP status from PubChem (e.g. 400 bad name).
        logger.error(f"PubChem HTTP error for '{drug_name}': {e}")
        return {"success": False,
                "error": f"Drug '{drug_name}' could not be found in PubChem — please check the spelling."}
    except (requests.exceptions.ConnectionError, requests.exceptions.Timeout) as e:
        # Genuine connectivity/timeout problem — do NOT cache this.
        logger.error(f"PubChem connection error for '{drug_name}': {e}")
        return {"success": False,
                "error": "PubChem is temporarily unreachable. Please try again in a moment."}
    except Exception as e:
        logger.error(f"PubChem API error for '{drug_name}': {e}")
        return {"success": False, "error": "PubChem lookup failed. Please try again."}

@retry(stop=stop_after_attempt(Config.MAX_RETRIES), wait=wait_exponential(multiplier=1, min=2, max=10))
def _pubmed_search_with_retry(term, max_results):
    """PubMed search with retry logic."""
    ncbi_rate_limiter.wait()
    
    esearch_url = f"{Config.NCBI_API_BASE}/esearch.fcgi"
    params = {
        "db": "pubmed",
        "term": term,
        "retmode": "json",
        "retmax": max_results
    }
    
    if Config.NCBI_API_KEY:
        params["api_key"] = Config.NCBI_API_KEY
    
    response = requests.get(esearch_url, params=params, timeout=Config.REQUEST_TIMEOUT)
    response.raise_for_status()
    
    data = response.json()
    return data.get("esearchresult", {}).get("idlist", [])

def search_pubmed_articles(group, max_results=5):
    """Search PubMed for relevant articles with caching."""
    cache_key = {
        'drug': external_term(group),
        'strain': group.get("strain", ""),
        'target': group.get("target_organ", ""),
        'max': max_results
    }
    
    cached = api_cache.get('pubmed_search', cache_key)
    if cached:
        return cached
    
    parts_full = [
        external_term(group),
        group.get("strain", ""),
        group.get("target_organ", ""),
        "mouse[Title/Abstract] OR mice[Title/Abstract]"
    ]
    term_full = " ".join([p for p in parts_full if p]).strip()
    
    if not term_full:
        return []
    
    try:
        pmids = _pubmed_search_with_retry(term_full, max_results)
        
        if not pmids:
            api_cache.set('pubmed_search', cache_key, [])
            return []
        
        ncbi_rate_limiter.wait()
        esummary_url = f"{Config.NCBI_API_BASE}/esummary.fcgi"
        params = {
            "db": "pubmed",
            "id": ",".join(pmids),
            "retmode": "json"
        }
        
        if Config.NCBI_API_KEY:
            params["api_key"] = Config.NCBI_API_KEY
        
        r2 = requests.get(esummary_url, params=params, timeout=Config.REQUEST_TIMEOUT)
        r2.raise_for_status()
        data2 = r2.json()
        
        result_block = data2.get("result", {})
        articles = []
        
        for pmid in pmids:
            raw = result_block.get(pmid)
            if not raw:
                continue
            
            pubdate = raw.get("pubdate", "") or raw.get("sortpubdate", "")
            year = pubdate.split(" ")[0] if pubdate else None
            
            articles.append({
                "pmid": pmid,
                "title": raw.get("title", ""),
                "journal": raw.get("fulljournalname", ""),
                "year": year,
                "pubmed_url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
                "source": "PubMed"
            })
        
        api_cache.set('pubmed_search', cache_key, articles)
        return articles
        
    except Exception as e:
        logger.error(f"PubMed search failed: {e}")
        return []

def search_europe_pmc(group, max_results=12):
    """Search Europe PMC, prioritising OPEN-ACCESS full-text papers and pulling a
    mix of recent and older (classic) work so the reference list isn't all new."""
    drug = external_term(group)
    parts = [drug, group.get("target_organ", ""), "toxicity", "mouse OR mice OR rat"]
    base_q = " ".join([p for p in parts if p]).strip()
    if not base_q:
        return []

    cache_key = {'q': base_q, 'max': max_results}
    cached = api_cache.get('europe_pmc', cache_key)
    if cached:
        return cached

    # Two open-access passes: newest first, then most-cited (surfaces older classics)
    queries = [
        (f"({base_q}) AND (OPEN_ACCESS:Y)", "P_PDATE_D desc"),
        (f"({base_q}) AND (OPEN_ACCESS:Y)", "CITED desc"),
    ]
    out, seen = [], set()
    try:
        for q, sort in queries:
            params = {"query": q, "format": "json", "pageSize": max_results, "sort": sort}
            r = requests.get(f"{Config.EUROPE_PMC_BASE}/search", params=params,
                             timeout=Config.REQUEST_TIMEOUT)
            r.raise_for_status()
            for rec in (r.json().get("resultList", {}).get("result", []) or []):
                title = (rec.get("title") or "").strip()
                if not title or title.lower() in seen:
                    continue
                seen.add(title.lower())
                pmcid = rec.get("pmcid") or ""
                if pmcid:                       # direct full-text (open access)
                    link = f"https://www.ncbi.nlm.nih.gov/pmc/articles/{pmcid}/"
                else:
                    link = f"https://europepmc.org/article/{rec.get('source', 'MED')}/{rec.get('id', '')}"
                out.append({
                    "title": title,
                    "authors": rec.get("authorString", ""),
                    "year": rec.get("pubYear"),
                    "venue": rec.get("journalTitle", ""),
                    "pmid": rec.get("pmid", ""),
                    "doi": rec.get("doi", ""),
                    "url": link,
                    "open_access": True,
                    "citations": rec.get("citedByCount", 0),
                    "source": "Europe PMC",
                })
        api_cache.set('europe_pmc', cache_key, out)
        return out
    except Exception as e:
        logger.error(f"Europe PMC search failed: {e}")
        return out


# Phrases that state animal loss, and phrases that state its absence. Polarity
# lives in the phrase itself, so the two searches stay separable without having
# to parse negation out of free text.
_MORTALITY_POS = ['"animals died"', '"mice died"', '"rats died"',
                  '"treatment-related mortality"', '"treatment related mortality"',
                  '"died during the study"']
_MORTALITY_NEG = ['"no mortality"', '"no animals died"', '"no mice died"',
                  '"no rats died"', '"no treatment-related mortality"',
                  '"all animals survived"']
# A positive hit must not actually be one of the negated forms. Europe PMC has
# no sentence-level context, so the negated phrasings are excluded by query.
_MORTALITY_NOT = ['"no animals died"', '"none of the animals died"',
                  '"no mice died"', '"no rats died"', '"no animals had died"']


def _epmc_count_and_cite(query, max_cites=3):
    """Hit count plus a few citable records for one Europe PMC query.

    Relevance order, not citation count: sorting by citations surfaces broad
    review articles that merely mention the terms, which are useless to a
    researcher trying to read what a comparable study actually observed.
    """
    params = {"query": query, "format": "json", "pageSize": max_cites}
    r = requests.get(f"{Config.EUROPE_PMC_BASE}/search", params=params,
                     timeout=Config.REQUEST_TIMEOUT)
    r.raise_for_status()
    data = r.json()
    cites = []
    for rec in (data.get("resultList", {}).get("result", []) or []):
        title = (rec.get("title") or "").strip()
        if not title:
            continue
        cites.append({
            "title": title,
            "year": rec.get("pubYear", ""),
            "journal": rec.get("journalTitle", ""),
            "url": f"https://europepmc.org/article/{rec.get('source', 'MED')}/{rec.get('id', '')}",
        })
    return int(data.get("hitCount", 0)), cites


def search_mortality_evidence(group):
    """What the published literature actually reports about animal loss.

    This is deliberately NOT an attrition prediction. Measured across Europe
    PMC, fewer than 0.5% of open-access rodent toxicity papers state animal
    deaths at all, the phrasing is inconsistent, and the papers that do report
    deaths are mostly acute LD50 studies designed to be lethal — so any rate
    learned from that corpus would describe the wrong population. What the
    corpus CAN support is a citable statement of what comparable studies
    reported, which the researcher then reads and judges.

    Returns a dict with counts, citations, and an explicit `found` flag, or a
    `found: False` result the interface is expected to state plainly.
    """
    term = external_term(group)          # compound class when confidential
    if not term:
        return {"found": False, "reason": "No compound term to search."}

    species = (group.get("species") or "").lower()
    animal = "rat OR rats" if "rat" in species else "mouse OR mice"
    # The compound must be what the paper is ABOUT (title/abstract), and reviews
    # are excluded — otherwise the hits are broad articles that merely mention
    # the drug somewhere in the body, which tell the researcher nothing.
    scope = (f'(TITLE:"{term}" OR ABSTRACT:"{term}") AND ({animal}) '
             f'AND (OPEN_ACCESS:Y) AND (HAS_FT:Y) NOT (PUB_TYPE:"Review")')

    cache_key = {"term": term.lower(), "animal": animal}
    cached = api_cache.get("mortality_evidence", cache_key)
    if cached:
        return cached

    def section_any(phrases):
        return " OR ".join(f"{sec}:{p}" for p in phrases for sec in ("RESULTS", "METHODS"))

    pos_q = (f'{scope} AND ({section_any(_MORTALITY_POS)})'
             + "".join(f' NOT (RESULTS:{p} OR METHODS:{p})' for p in _MORTALITY_NOT))
    neg_q = f'{scope} AND ({section_any(_MORTALITY_NEG)})'

    try:
        n_pos, cites_pos = _epmc_count_and_cite(pos_q)
        n_neg, cites_neg = _epmc_count_and_cite(neg_q)
    except Exception as e:
        logger.warning(f"Mortality-evidence search failed for {term}: {e}")
        return {"found": False, "reason": "Literature search unavailable."}

    result = {
        "found": bool(n_pos or n_neg),
        "term": term,
        "reported_mortality": n_pos,
        "reported_no_mortality": n_neg,
        "citations_mortality": cites_pos,
        "citations_no_mortality": cites_neg,
        "caveat": ("Counts of what open-access papers state, not an attrition rate. "
                   "Reporting is sparse and inconsistent, and studies reporting deaths "
                   "are often acute lethality designs — read the cited papers before "
                   "relying on them."),
    }
    if not result["found"]:
        result["reason"] = "No open-access study of this compound states either outcome."
    api_cache.set("mortality_evidence", cache_key, result)
    return result


def search_doaj(group, max_results=6):
    """Search DOAJ (Directory of Open Access Journals) — fully open-access,
    peer-reviewed journal articles."""
    parts = [external_term(group), group.get("target_organ", ""), "toxicity"]
    q = " ".join([p for p in parts if p]).strip()
    if not q:
        return []
    cache_key = {'q': q, 'max': max_results}
    cached = api_cache.get('doaj', cache_key)
    if cached:
        return cached
    try:
        from urllib.parse import quote
        url = f"{Config.DOAJ_BASE}/{quote(q)}"
        r = requests.get(url, params={"pageSize": max_results, "sort": "created_date:desc"},
                         timeout=Config.REQUEST_TIMEOUT)
        r.raise_for_status()
        out = []
        for res in (r.json().get("results", []) or []):
            b = res.get("bibjson", {}) or {}
            doi = ""
            link = ""
            for idf in b.get("identifier", []) or []:
                if idf.get("type") == "doi":
                    doi = idf.get("id", "")
            for lk in b.get("link", []) or []:
                if lk.get("url"):
                    link = lk["url"]
                    break
            authors = ", ".join(a.get("name", "") for a in (b.get("author", []) or [])[:6])
            out.append({
                "title": b.get("title", ""),
                "authors": authors,
                "year": b.get("year"),
                "venue": (b.get("journal", {}) or {}).get("title", ""),
                "doi": doi,
                "url": link or (f"https://doi.org/{doi}" if doi else ""),
                "open_access": True,
                "source": "DOAJ",
            })
        api_cache.set('doaj', cache_key, out)
        return out
    except Exception as e:
        logger.error(f"DOAJ search failed: {e}")
        return []


def search_preprints(group, max_results=6):
    """Search bioRxiv / medRxiv preprints (indexed and served via Europe PMC)."""
    parts = [external_term(group), group.get("target_organ", ""), "toxicity",
             "mouse OR mice OR rat"]
    base_q = " ".join([p for p in parts if p]).strip()
    if not base_q:
        return []
    cache_key = {'q': base_q, 'max': max_results}
    cached = api_cache.get('preprints', cache_key)
    if cached:
        return cached
    try:
        q = f'({base_q}) AND (SRC:"PPR")'   # PPR = preprint sources (bioRxiv/medRxiv…)
        r = requests.get(f"{Config.EUROPE_PMC_BASE}/search",
                         params={"query": q, "format": "json", "pageSize": max_results,
                                 "sort": "P_PDATE_D desc"},
                         timeout=Config.REQUEST_TIMEOUT)
        r.raise_for_status()
        out = []
        for rec in (r.json().get("resultList", {}).get("result", []) or []):
            title = (rec.get("title") or "").strip()
            if not title:
                continue
            out.append({
                "title": title,
                "authors": rec.get("authorString", ""),
                "year": rec.get("pubYear"),
                "venue": rec.get("bookOrReportDetails", {}).get("publisher", "") or "Preprint (bioRxiv/medRxiv)",
                "doi": rec.get("doi", ""),
                "url": f"https://europepmc.org/article/{rec.get('source', 'PPR')}/{rec.get('id', '')}",
                "open_access": True,
                "source": "Preprint (bioRxiv/medRxiv)",
            })
        api_cache.set('preprints', cache_key, out)
        return out
    except Exception as e:
        logger.error(f"Preprint search failed: {e}")
        return []


def search_clinical_trials(drug, max_results=3, timeout=4):
    """ClinicalTrials.gov (free, no API key) — trials for a given drug.

    Uses a STRICT short timeout and swallows every error so it can never hang or
    break the results. Returns {total, studies:[...]} or None on any failure.
    """
    drug = (drug or "").strip()
    if not drug:
        return None
    cache_key = {'drug': drug.lower(), 'max': max_results}
    cached = api_cache.get('clinicaltrials', cache_key)
    if cached is not None:
        return cached
    try:
        r = requests.get(
            "https://clinicaltrials.gov/api/v2/studies",
            params={'query.intr': drug, 'countTotal': 'true', 'pageSize': max_results,
                    'fields': 'NCTId,BriefTitle,OverallStatus,Phase'},
            timeout=timeout)
        if r.status_code != 200:
            return {'total': 0, 'studies': []}
        data = r.json()
        studies = []
        for s in (data.get('studies') or [])[:max_results]:
            p = s.get('protocolSection', {}) or {}
            idm = p.get('identificationModule', {}) or {}
            stat = p.get('statusModule', {}) or {}
            des = p.get('designModule', {}) or {}
            nct = idm.get('nctId', '')
            studies.append({
                'nct': nct,
                'title': idm.get('briefTitle', ''),
                'status': stat.get('overallStatus', ''),
                'phase': ', '.join(des.get('phases', []) or []),
                'url': f"https://clinicaltrials.gov/study/{nct}" if nct else '',
            })
        result = {'total': int(data.get('totalCount', 0) or 0), 'studies': studies}
        api_cache.set('clinicaltrials', cache_key, result)
        return result
    except Exception as e:
        logger.warning(f"ClinicalTrials.gov lookup failed/timed out for {drug}: {e}")
        return None


def search_impc(group, max_results=3):
    """Search IMPC for strain phenotype data."""
    strain = group.get('strain', '').replace('/', '%2F')
    target = group.get('target_organ', '').lower()
    
    phenotype_map = {
        'brain': ['nervous system', 'behavior'],
        'heart': ['cardiovascular system'],
        'liver': ['liver', 'metabolism'],
        'kidney': ['renal'],
        'lung': ['respiratory system'],
    }
    
    phenotypes = []
    for key, terms in phenotype_map.items():
        if key in target:
            phenotypes = terms
            break
    
    if not strain or not phenotypes:
        return []
    
    cache_key = {'strain': strain, 'phenotypes': str(phenotypes), 'max': max_results}
    cached = api_cache.get('impc', cache_key)
    if cached:
        return cached
    
    try:
        base_url = f"{Config.IMPC_BASE}/genotype-phenotype/select"
        results = []
        
        for phenotype in phenotypes[:2]:
            params = {
                'q': f'marker_symbol:* AND strain_name:*{strain}* AND top_level_mp_term_name:*{phenotype}*',
                'rows': max_results,
                'wt': 'json'
            }
            
            response = requests.get(base_url, params=params, timeout=Config.REQUEST_TIMEOUT)
            response.raise_for_status()
            data = response.json()
            
            docs = data.get('response', {}).get('docs', [])
            for doc in docs:
                results.append({
                    'strain': doc.get('strain_name'),
                    'gene': doc.get('marker_symbol'),
                    'phenotype': doc.get('mp_term_name'),
                    'parameter': doc.get('parameter_name'),
                    'p_value': doc.get('p_value'),
                    'effect_size': doc.get('effect_size'),
                    'url': f"https://www.mousephenotype.org/data/genes/{doc.get('marker_accession_id', '')}",
                    'source': 'IMPC'
                })
        
        results = results[:max_results]
        api_cache.set('impc', cache_key, results)
        return results
        
    except Exception as e:
        logger.error(f"IMPC search failed: {e}")
        return []

def build_comprehensive_reference_corpus(group):
    """Build comprehensive reference corpus from ALL sources."""
    logger.info("Building comprehensive reference corpus for "
                + (external_term(group) or "unnamed compound"))
    
    # Search all APIs CONCURRENTLY so total time ≈ the slowest single call
    # (not the sum). One slow/failing database no longer stalls the request.
    from concurrent.futures import ThreadPoolExecutor
    _tasks = {
        'pubmed':   lambda: search_pubmed_articles(group, max_results=6),
        'europe':   lambda: search_europe_pmc(group, max_results=12),
        'semantic': lambda: search_semantic_scholar(group, max_results=6),
        'openalex': lambda: search_openalex(group, max_results=6),
        'crossref': lambda: search_crossref(group, max_results=4),
        'doaj':     lambda: search_doaj(group, max_results=6),
        'preprint': lambda: search_preprints(group, max_results=6),
        'impc':     lambda: search_impc(group, max_results=3),
    }
    _results = {k: [] for k in _tasks}
    # Start every source at once; cap each result-wait; and DON'T block the
    # request waiting for a slow/hung source (e.g. a rate-limited retry) to
    # finish — shutdown(wait=False) lets stragglers die in the background.
    _ex = ThreadPoolExecutor(max_workers=len(_tasks))
    _futs = {_ex.submit(fn): name for name, fn in _tasks.items()}
    for _fut, _name in _futs.items():
        try:
            _results[_name] = _fut.result(timeout=12) or []
        except Exception as e:
            logger.warning(f"{_name} search failed/timeout: {e}")
            _results[_name] = []
    _ex.shutdown(wait=False)
    pubmed_refs = _results['pubmed']
    europe_refs = _results['europe']
    semantic_refs = _results['semantic']
    openalex_refs = _results['openalex']
    crossref_refs = _results['crossref']
    doaj_refs = _results['doaj']
    preprint_refs = _results['preprint']
    impc_refs = _results['impc']

    # Combine and deduplicate by title (open-access sources included)
    all_papers = []
    seen_titles = set()

    for paper_list in [europe_refs, doaj_refs, pubmed_refs, semantic_refs,
                       openalex_refs, crossref_refs, preprint_refs]:
        for paper in paper_list:
            title = paper.get('title', '').lower().strip()
            if title and title not in seen_titles:
                seen_titles.add(title)
                all_papers.append(paper)
    
    def safe_int(val):
        try:
            return int(val)
        except (TypeError, ValueError):
            return 0

    def _is_oa(p):
        if p.get('open_access'):
            return True
        u = (p.get('url') or '').lower()
        return 'europepmc.org' in u or '/pmc/' in u or 'ncbi.nlm.nih.gov/pmc' in u

    # Build a list that is (a) open-access first and (b) a spread of years — by
    # interleaving the newest and the oldest OA papers — so the references are
    # never all-recent and always include freely readable full-text sources.
    oa = [p for p in all_papers if _is_oa(p)]
    other = [p for p in all_papers if not _is_oa(p)]
    oa_new = sorted(oa, key=lambda x: safe_int(x.get('year')), reverse=True)
    oa_old = sorted(oa, key=lambda x: safe_int(x.get('year')))
    picked, seen_ids = [], set()

    def _add(p):
        if id(p) not in seen_ids:
            seen_ids.add(id(p))
            picked.append(p)

    # Guarantee visibility for the newer open sources (DOAJ journals + preprints)
    # so the reference list isn't only Europe PMC.
    for _src in (doaj_refs, preprint_refs):
        _added = 0
        for p in _src:
            if _added >= 3:
                break
            _add(p); _added += 1

    i = j = 0
    while len(picked) < 20 and (i < len(oa_new) or j < len(oa_old)):
        if i < len(oa_new):
            _add(oa_new[i]); i += 1
        if j < len(oa_old):
            _add(oa_old[j]); j += 1
    for p in sorted(other, key=lambda x: safe_int(x.get('citations')), reverse=True):
        if len(picked) >= 20:
            break
        _add(p)

    return {
        "pubmed": pubmed_refs,
        "europe_pmc": europe_refs,
        "semantic_scholar": semantic_refs,
        "openalex": openalex_refs,
        "crossref": crossref_refs,
        "doaj": doaj_refs,
        "preprint": preprint_refs,
        "impc": impc_refs,
        "all_papers": picked[:20]
    }

# ============================================================================
# STATISTICAL FUNCTIONS
# ============================================================================

# ============================================================================
# OUTCOME VARIABILITY MODEL (OVM)
# ----------------------------------------------------------------------------
# Sample size is driven by how much the measured outcome varies between animals.
# A fixed Cohen's d = 0.8 returns the same N for every experiment, which is
# wrong in both directions: body weight varies ~11% between mice while operant
# behavioural counts vary ~100%. This predicts the coefficient of variation for
# the specific measure/strain/sex so the required N follows the real experiment.
#
# Trained on curated phenotype databases (Mouse Phenome Database + RGD
# PhenoMiner). It beats a per-measure median by ~11% on unseen strain/sex cells
# but adds nothing for a measure absent from both — so an absent measure gets no
# prediction and the caller says so, rather than guessing.
# ============================================================================
_OVM = None
_OVM_META = None
_OVM_MEASURES = None


def _load_ovm():
    """Load the variability model once, tolerating its absence."""
    global _OVM, _OVM_META, _OVM_MEASURES
    if _OVM is not None or _OVM is False:
        return
    try:
        _OVM = joblib.load(os.path.join(Config.ML_MODEL_PATH, 'variability_model.pkl'))
        _OVM_META = json.load(open(os.path.join(Config.ML_MODEL_PATH,
                                                'variability_meta.json')))
        _OVM_MEASURES = json.load(open(os.path.join(Config.ML_MODEL_PATH,
                                                    'variability_measures.json')))
        logger.info("Loaded Outcome Variability Model (%s rows, %s measures)",
                    _OVM_META.get('n_rows'), _OVM_META.get('n_measures'))
    except Exception as e:
        _OVM = False
        logger.warning(f"Outcome Variability Model unavailable: {e}")


def predict_outcome_cv(measure, species, strain=None, sex=None):
    """Predicted coefficient of variation (%) for one outcome measure.

    Returns None when the measure is outside the training databases — held-out
    testing showed the model adds nothing there, so reporting a number would be
    false precision on the one figure an ethics committee scrutinises most.
    """
    _load_ovm()
    if not _OVM or not measure:
        return None
    key = ((species or 'mouse').strip().lower(), measure.strip().lower())
    if key not in _OVM['per_measure']:
        return {'in_domain': False, 'measure': measure, 'species': key[0],
                'reason': 'This outcome measure is not in the reference databases.'}
    try:
        import numpy as np
        vocab, encoders = _OVM['encoders']
        cats, base = _OVM['cats'], _OVM['per_measure'][key]
        row = {'species': key[0], 'measure': key[1],
               'strain': (strain or 'unspecified'), 'sex': (sex or 'unspecified'),
               'intervention': '', 'method': ''}
        x = np.zeros((1, 1 + len(cats) + 2), dtype=float)
        x[0, 0] = np.log10(base)
        log_global = np.log10(_OVM['global_median'])
        for j, c in enumerate(cats, start=1):
            v = row.get(c, '') or ''
            x[0, j] = vocab[c].get(v, -1) if c in vocab else encoders[c].get(v, log_global)
        x[0, len(cats) + 1] = 8          # typical group size
        x[0, len(cats) + 2] = 0.0
        cv = float(10 ** _OVM['model'].predict(x)[0])
        n_obs = next((m['n_observations'] for m in (_OVM_MEASURES or [])
                      if m['species'] == key[0] and m['measure'] == key[1]), None)
        return {
            'in_domain': True, 'measure': measure, 'species': key[0],
            'cv_pct': round(cv, 1), 'measure_median_cv_pct': round(base, 1),
            'n_reference_observations': n_obs,
            'model': (_OVM_META or {}).get('name', 'Outcome Variability Model'),
            'test_mae_cv_points': round((_OVM_META or {}).get(
                'test_mae_cell_held_out', 0), 1),
            'sources': (_OVM_META or {}).get('sources', []),
        }
    except Exception as e:
        logger.warning(f"Variability prediction failed for {measure}: {e}")
        return None


def calculate_sample_size_power_analysis(
    effect_size: float,
    alpha: float = 0.05,
    power: float = 0.80,
    sd: float = None,
    mean_diff: float = None
) -> dict:
    """Statistical power analysis for sample size determination."""
    try:
        if mean_diff is not None and sd is not None and sd > 0:
            effect_size = abs(mean_diff) / sd
        
        if effect_size <= 0:
            effect_size = 0.5
        
        n_per_group = tt_ind_solve_power(
            effect_size=effect_size,
            alpha=alpha,
            power=power,
            ratio=1.0,
            alternative='two-sided'
        )
        
        n_per_group_rounded = int(np.ceil(n_per_group))
        n_with_attrition = int(np.ceil(n_per_group_rounded * 1.10))
        
        achieved_power = tt_ind_solve_power(
            effect_size=effect_size,
            alpha=alpha,
            nobs1=n_per_group_rounded,
            ratio=1.0,
            alternative='two-sided'
        )
        
        return {
            "success": True,
            "n_per_group": n_per_group_rounded,
            "n_with_10pct_attrition": n_with_attrition,
            "requested_power": power,
            "achieved_power": round(achieved_power, 3),
            "alpha": alpha,
            "effect_size": round(effect_size, 3),
            "method": "Two-sample t-test power analysis"
        }
    except Exception as e:
        logger.error(f"Power analysis failed: {e}")
        return {
            "success": False,
            "error": str(e),
            "n_per_group": 10
        }

def generate_power_curve_data(effect_size, alpha=0.05):
    """Generate data for power curve visualization."""
    sample_sizes = np.arange(5, 31, 1)
    powers = []
    
    for n in sample_sizes:
        try:
            power = tt_ind_solve_power(
                effect_size=effect_size,
                alpha=alpha,
                nobs1=n,
                ratio=1.0,
                alternative='two-sided'
            )
            powers.append(power)
        except:
            powers.append(0.5)
    
    return sample_sizes.tolist(), powers

# ============================================================================
# SAMPLE TYPES & UTILITIES
# ============================================================================

SAMPLE_TYPES = [
    "Blood (whole blood)", "Blood (plasma)", "Blood (serum)",
    "Brain tissue", "Heart tissue", "Liver tissue", "Kidney tissue",
    "Lung tissue", "Spleen tissue", "Muscle tissue", "Bone tissue",
    "Bone marrow", "Adipose tissue", "Pancreatic tissue",
    "Intestinal tissue", "Skin tissue", "Tumor tissue",
    "Cerebrospinal fluid (CSF)", "Urine", "Feces", "Other"
]

def parse_float_safe(value, default=0.0):
    """Safely parse float values."""
    try:
        return float(value)
    except (TypeError, ValueError):
        return default

def parse_int_safe(value, default=0):
    """Safely parse integer values."""
    try:
        return int(value)
    except (TypeError, ValueError):
        return default

def recommend_statistical_test(num_groups: int, repeated_measures: bool = False) -> dict:
    """Recommend appropriate statistical test."""
    recommendations = {
        "test": "",
        "assumptions": [],
        "software": [],
        "post_hoc": ""
    }
    
    if num_groups == 2:
        if repeated_measures:
            recommendations["test"] = "Paired t-test"
            recommendations["assumptions"] = [
                "Paired observations",
                "Normally distributed differences"
            ]
            recommendations["software"] = ["R: t.test(paired=TRUE)", "Python: scipy.stats.ttest_rel"]
        else:
            recommendations["test"] = "Independent samples t-test"
            recommendations["assumptions"] = [
                "Independent groups",
                "Normal distribution",
                "Equal variances"
            ]
            recommendations["software"] = ["R: t.test()", "Python: scipy.stats.ttest_ind"]
    else:
        recommendations["test"] = "One-way ANOVA"
        recommendations["assumptions"] = [
            "Independent groups",
            "Normal distribution",
            "Homogeneity of variances"
        ]
        recommendations["software"] = ["R: aov()", "Python: scipy.stats.f_oneway"]
        recommendations["post_hoc"] = "Tukey HSD or Bonferroni"
    
    return recommendations


def get_compound_summary(drug_name, cid=None):
    """A brief plain-language description of the compound from PubChem (cached),
    with a citable reference URL. Used to give the researcher a quick overview."""
    ck = {'drug': (drug_name or '').lower(), 'type': 'desc'}
    cached = api_cache.get('pubchem_desc', ck)
    if cached:
        return cached
    result = {'description': '', 'reference_url': ''}
    try:
        if cid is None:
            cd = get_drug_data_from_ncbi(drug_name)
            cid = cd.get('cid') if cd.get('success') else None
        if not cid:
            return result
        result['reference_url'] = f"https://pubchem.ncbi.nlm.nih.gov/compound/{cid}"
        r = requests.get(f"{Config.PUBCHEM_API_BASE}/compound/cid/{cid}/description/JSON", timeout=8)
        if r.status_code == 200:
            for info in r.json().get('InformationList', {}).get('Information', []):
                if info.get('Description'):
                    result['description'] = info['Description']
                    if info.get('DescriptionURL'):
                        result['reference_url'] = info['DescriptionURL']
                    break
    except Exception as e:
        logger.warning(f"Compound description failed for {drug_name}: {e}")
    if result['description']:
        api_cache.set('pubchem_desc', ck, result)
    return result


def build_protocol_timeline(group, animal_word='mice'):
    """
    Generate a day-by-day experimental protocol as an editable STARTING POINT
    (acclimatize -> baseline -> fast -> dose -> monitor -> sample -> endpoint).
    Concrete day numbers are defaults the researcher can edit.
    """
    dose = group.get('dose', '')
    route = (group.get('route', 'Oral') or 'Oral')
    diet = (group.get('diet_type', 'Standard') or 'Standard')
    samples = group.get('sample_types', []) or []
    sample_str = ', '.join(samples) if samples else 'blood/tissue as required'
    is_control = (group.get('drug_name', '') or '').lower() in ['saline', 'control', 'pbs', 'vehicle']
    agent = 'vehicle (saline)' if is_control else f"{group.get('drug_name', 'compound')} at {dose} mg/kg"

    # Group-specific objective — so Sample collection & Endpoint rows read
    # per group (each group's aims differ within the study).
    _organ = (group.get('target_organ') or '').strip()
    _endpoints = group.get('toxicity_endpoints', []) or []
    _obj_bits = []
    if _endpoints:
        _obj_bits.append(', '.join(_endpoints))
    if _organ and _organ.lower() not in ('', 'general', 'none'):
        _obj_bits.append(f'target organ: {_organ}')
    if is_control:
        objective = 'baseline / control comparison'
    else:
        objective = '; '.join(_obj_bits) if _obj_bits else 'general toxicity assessment'

    steps = [
        {'day': 'Day 1–3', 'phase': 'Acclimatization',
         'activity': f'Allow {animal_word} to acclimatize to the facility. Free access to food and water; daily health checks.'},
        {'day': 'Day 4', 'phase': 'Baseline',
         'activity': 'Record baseline body weight and body condition score; randomize animals into groups.'},
    ]
    if diet.lower() == 'fast':
        steps.append({'day': 'Day 5 (AM)', 'phase': 'Fasting',
                      'activity': 'Fast animals 6–12 h before dosing (water allowed); record pre-dose weight.'})
    steps += [
        {'day': 'Day 5', 'phase': 'Dosing',
         'activity': f'Administer {agent} via {route} route; record time and any immediate reactions.'},
        {'day': 'Day 6–11', 'phase': 'Monitoring',
         'activity': 'Daily monitoring: body weight, clinical signs and humane endpoints (see Toxicity & Humane Endpoints).'},
        {'day': 'Day 12', 'phase': 'Sample collection',
         'activity': f'Collect samples: {sample_str} — for {objective}. Respect safe blood-volume limits.'},
        {'day': 'Day 13', 'phase': 'Endpoint',
         'activity': f'Apply humane endpoint / approved euthanasia; collect terminal tissues for {objective}.'},
    ]
    return steps


def biological_advice_for(species, group):
    """Short species-specific husbandry / handling advice."""
    common = [
        'Confirm veterinary consultation and IACUC approval before starting.',
        'Use aseptic technique and appropriate analgesia where procedures may cause pain.',
    ]
    if species == 'Rat':
        return [
            'Rats: house in pairs/groups; provide gnawing/enrichment. ~200–350 g adult.',
            'Oral gavage volume ≤ 10 mL/kg; handle calmly to reduce stress.',
        ] + common
    return [
        'Mice: house in social groups (≤5/cage); provide nesting material. ~20–30 g adult.',
        'Oral gavage volume ≤ 10 mL/kg; minimize handling stress.',
    ] + common


# ============================================================================
# CORE PREDICTION LOGIC
# ============================================================================

def get_prediction_and_suggestion(group, all_groups_count=1):
    """Main prediction logic with ML and comprehensive literature search."""
    drug = group.get('drug_name', '') or ''
    is_control = drug.lower() in ['saline', 'control', 'pbs', 'vehicle']

    # Species-aware wording (Mouse vs Rat) used throughout the output
    species = (group.get('species') or 'Mouse').strip().title()
    if species not in ('Mouse', 'Rat'):
        species = 'Mouse'
    animal_word = 'rats' if species == 'Rat' else 'mice'
    # Standard husbandry reference ranges per species (adult)
    STD_REF = {
        'Mouse': {'weight': '20-30 g', 'age': '6-12 weeks'},
        'Rat':   {'weight': '200-350 g', 'age': '8-14 weeks'},
    }[species]
    
    # Build comprehensive reference corpus from ALL sources
    ref_corpus = build_comprehensive_reference_corpus(group)
    has_refs = len(ref_corpus.get("all_papers", [])) > 0
    
    # Calculate blood quantity if blood samples selected
    blood_calc = BloodQuantityCalculator.calculate_blood_needed(
        sample_types=group.get('sample_types', []),
        weight_g=parse_float_safe(group.get('weight'), 25),
        timepoints=max(1, int(parse_float_safe(group.get('blood_timepoints'), 1))),
        num_replicates=max(1, int(parse_float_safe(group.get('blood_replicates'), 1)))
    )

    _private = is_confidential(group)

    def build_summary(recommended_range):
        """A clear, simple recommendation summary for the UI."""
        # GHS acute-toxicity class from the LD50 (experimental first, then model
        # estimate). No invented percentage — only the published GHS band.
        ghs = tox_cat = ld50_val = tox_source = None
        if not is_control:
            try:
                if _private:
                    # Name-based lookup would expose the compound — predict from
                    # the pasted structure with the local model instead.
                    info = {'found': False}
                    _ml = toxicity_predictor.predict_ld50_ml((group.get('smiles') or '').strip())
                    if _ml:
                        info = {'found': True, 'ld50_mg_kg': _ml['ld50_mg_kg'],
                                'category': _ml['category'], 'source': 'ml_model'}
                else:
                    info = get_real_ld50_info(drug, route=group.get('route', 'oral'), species=species)
                if info.get('found'):
                    ld50_val = info['ld50_mg_kg']
                    tox_cat = info['category']
                    ghs = ld50_to_ghs(ld50_val)
                    tox_source = info.get('source')  # 'experimental' or 'ml_model'
                    # An LD50 the model inferred carries a wide error band
                    # (test MAE 0.48 log units ~ a factor of 3), so it is
                    # reported to 2 significant figures with that band stated.
                    if tox_source == 'ml_model':
                        ld50_val = round_sig(ld50_val, 2)
                    else:
                        ld50_val = round_sig(ld50_val, 3)
            except Exception as e:
                logger.warning(f"Summary toxicity lookup failed for {drug}: {e}")

        # Humane endpoints + pain monitoring (IACUC Parts 8 & 11), derived from
        # the toxicity level and dose (for treatment groups only).
        welfare = None
        if not is_control:
            try:
                dose_val = parse_float_safe(group.get('dose'), 0)
                pct_ld50 = (dose_val / ld50_val * 100) if (ld50_val and dose_val) else 0
                # With no toxicity data we fall back to a generic moderate
                # setting; `evidence_based` tells the UI to say so plainly
                # rather than let a placeholder pass as a finding.
                risk_proxy = ghs['risk_score'] if ghs else 40
                welfare = toxicity_predictor.generate_welfare_recommendations(
                    risk_proxy, pct_ld50=pct_ld50,
                    target_organ=group.get('target_organ'))
                if welfare is not None:
                    welfare['evidence_based'] = ghs is not None
            except Exception as e:
                logger.warning(f"Welfare recommendation failed for {drug}: {e}")
        # Recommended tissues/samples to collect, based on the chosen toxicity
        # endpoints and the target organ.
        _endpoints = group.get('toxicity_endpoints', []) or []
        _organ = (group.get('target_organ') or '').lower()
        rec_samples = []
        # Substring keyword -> recommended samples. Covers the categorized
        # Outcome Measures so a user's selection actually drives sampling.
        _ep_map = {
            # classic toxicity endpoints (custom entries may still use these)
            'hepatotox': ['Liver tissue', 'Blood (serum)'],
            'nephrotox': ['Kidney tissue', 'Blood (serum)'],
            'cardiotox': ['Heart tissue', 'Blood (plasma)'],
            'neurolog':  ['Brain tissue'],
            'histopath': ['Target-organ tissue'],
            # Histology
            'h&e': ['Target-organ tissue (for histology)'],
            'ihc': ['Target-organ tissue (for IHC)'],
            'oil red': ['Liver tissue', 'Adipose tissue'],
            # Cardiovascular
            'ecg': ['Heart tissue', 'Blood (plasma)'],
            'blood pressure': ['Blood (plasma)'],
            'lipid': ['Blood (serum)'],
            # Metabolic
            'glucose': ['Blood (serum)'],
            'ogtt': ['Blood (serum)'],
            'itt': ['Blood (serum)'],
            'insulin': ['Blood (serum)'],
            'homa': ['Blood (serum)'],
            # Inflammation
            'tnf': ['Blood (serum)', 'Spleen tissue'],
            'il-6': ['Blood (serum)', 'Spleen tissue'],
            'il-1': ['Blood (serum)', 'Spleen tissue'],
            'crp': ['Blood (serum)'],
            # Molecular
            'western blot': ['Target-organ tissue (for protein)'],
            'elisa': ['Blood (serum)'],
            'qpcr': ['Target-organ tissue (for RNA)'],
            # Oncology
            'tumor': ['Tumor tissue'],
            'metasta': ['Lung tissue', 'Liver tissue'],
            'ki-67': ['Tumor tissue'],
            # Vaccines / Infection
            'virus titer': ['Blood (serum)', 'Target-organ tissue'],
        }
        for ep in _endpoints:
            for kw, samps in _ep_map.items():
                if kw in ep.lower():
                    rec_samples += samps
        organ_tissue = {'liver': 'Liver tissue', 'kidney': 'Kidney tissue',
                        'heart': 'Heart tissue', 'brain': 'Brain tissue',
                        'lung': 'Lung tissue', 'spleen': 'Spleen tissue'}.get(_organ)
        if organ_tissue:
            rec_samples.append(organ_tissue)
        if not rec_samples:
            rec_samples = ['Blood (plasma)', 'Target-organ tissue']
        # de-duplicate, preserve order
        recommended_samples = list(dict.fromkeys(rec_samples))

        # Brief compound overview + reference (skip for controls)
        drug_overview = None
        if not is_control and not _private:      # both look the compound up by name
            try:
                drug_overview = get_compound_summary(drug)
            except Exception as e:
                logger.warning(f"Compound overview failed for {drug}: {e}")

        # ClinicalTrials.gov context (free, no key; strict 4s timeout — the
        # function swallows any error/timeout, so results never hang). Controls skipped.
        clinical_trials = None if (is_control or _private) else search_clinical_trials(drug)

        # What comparable published studies reported about animal loss. Safe in
        # confidential mode: the search runs on the compound CLASS, never the
        # name (search_mortality_evidence uses external_term).
        mortality_evidence = None if is_control else search_mortality_evidence(group)

        # Recommended strain by experimental paradigm (common choices in the
        # literature) — a suggestion the researcher can override.
        _paradigm = (group.get('experiment_type') or '').lower()
        _strain_rec = {
            'oncology': 'BALB/c nude or NOD-SCID (immunodeficient, for xenografts)',
            'immunology': 'C57BL/6 or BALB/c',
            'neuroscience': 'C57BL/6',
            'metabolic': 'C57BL/6 (diet-induced) or ob/ob',
            'cardiovascular': 'C57BL/6',
            'toxicology': 'C57BL/6 (mouse) or Sprague-Dawley (rat)',
        }
        recommended_strain = None
        for kw, rec in _strain_rec.items():
            if kw in _paradigm:
                recommended_strain = rec
                break
        if not recommended_strain:
            recommended_strain = 'Sprague-Dawley / Wistar' if species == 'Rat' else 'C57BL/6 (most common)'

        # ML property estimates from structure — inform vehicle/formulation
        # (solubility) and dosing frequency (half-life). Skip controls.
        solubility = None
        halflife = None
        ml_flags = []
        ml_adme = []
        if not is_control:
            try:
                # Confidential compound: use the structure the investigator pasted,
                # so the name is never sent to PubChem. Models run locally either way.
                if _private:
                    _smi = (group.get('smiles') or '').strip()
                    _struct = {'success': bool(_smi), 'smiles': _smi}
                else:
                    _struct = toxicity_predictor.get_chemical_structure(drug)
                if _struct.get('success'):
                    _smi = _struct['smiles']
                    _feats = toxicity_predictor._featurize_smiles(_smi)   # once, shared
                    solubility = toxicity_predictor.predict_solubility_ml(_smi, _feats)
                    halflife = toxicity_predictor.predict_halflife_ml(_smi, _feats)
                    ml_flags = toxicity_predictor.predict_flags_ml(_smi, _feats)
                    ml_adme = toxicity_predictor.predict_adme_ml(_smi, _feats)
            except Exception as e:
                logger.warning(f"ML property prediction failed for {drug}: {e}")

        # Toxicity-based safer-dosing guidance, derived from the estimated LD50
        # (a SAFETY margin, not an efficacy range). Skip controls.
        safe_dose = None
        if not is_control and ld50_val and tox_cat:
            try:
                guide = toxicity_predictor.get_safe_dose_recommendations(
                    tox_cat, group.get('route', 'oral'))
                dose_val = parse_float_safe(group.get('dose'), 0)
                safe_dose = {
                    'ld50_mg_kg': ld50_val,
                    'ld50_source': tox_source,          # experimental / ml_model
                    'starting_dose': guide.get('starting_dose'),
                    'max_dose': guide.get('max_dose'),
                    'safety_ceiling_mg_kg': round(ld50_val * 0.1, 1),   # ~1/10 LD50
                    'current_dose': dose_val or None,
                    'current_pct_ld50': round(dose_val / ld50_val * 100, 1) if dose_val else None,
                    'monitoring': guide.get('monitoring'),
                }
            except Exception as e:
                logger.warning(f"Safe-dose derivation failed for {drug}: {e}")

        return {
            'species': species,
            'strain': group.get('strain', ''),
            'recommended_strain': recommended_strain,
            'drug_overview': drug_overview,
            'animal_word': animal_word,
            'recommended_samples': recommended_samples,
            'recommended_animals': recommended_range,
            'planned_animals': group.get('num_mice', ''),
            'planned_weight_g': group.get('weight'),
            'planned_age_weeks': group.get('age'),
            'standard_weight': STD_REF['weight'],
            'standard_age': STD_REF['age'],
            'route': group.get('route', ''),
            'diet_type': group.get('diet_type', 'Standard'),
            'experiment_type': group.get('experiment_type', ''),
            'target_organ': group.get('target_organ', 'General'),
            'toxicity_endpoints': group.get('toxicity_endpoints', []),
            'samples': group.get('sample_types', []),
            'blood_volume_ml': blood_calc.get('total_volume_ml') if blood_calc.get('needed') else None,
            'blood_safe': (blood_calc.get('safety_color') == 'green') if blood_calc.get('needed') else None,
            'ghs': ghs,                     # published GHS acute-toxicity band
            'toxicity_category': tox_cat,
            'ld50_mg_kg': ld50_val,
            'toxicity_source': tox_source,  # experimental vs ml_model (estimate)
            'ld50_error_band': ('×/÷ 3 (model test MAE 0.48 log units)'
                                if tox_source == 'ml_model' else None),
            'solubility': solubility,       # ML aqueous-solubility estimate (vehicle hint)
            'halflife': halflife,           # ML half-life estimate (human PK; frequency hint)
            'safe_dose': safe_dose,         # LD50-derived safer-dosing guidance (safety margin)
            'ml_flags': ml_flags,           # ML safety/ADME classifier flags (hERG, DILI, Ames, BBB)
            'ml_adme': ml_adme,             # ML ADME regression values (LogD, Caco-2, PPBR)
            'clinical_trials': clinical_trials,  # ClinicalTrials.gov context (free)
            'mortality_evidence': mortality_evidence,  # cited, never a predicted rate
            'attrition_basis': ('Standard 10% allowance for technical loss. Not adjusted for '
                                'predicted toxicity — a dose expected to kill animals should be '
                                'lowered, not padded with extra animals.'),
            'is_control': is_control,       # lets the UI match control N to treatment N
            'confidential': _private,       # compound name kept inside the platform
            'timeline': build_protocol_timeline(group, animal_word),
            'welfare': welfare,
            'biological_advice': biological_advice_for(species, group),
        }

    if is_control:
        # The control gets the same power-analysis N (balanced design). The
        # number is NOT clamped: capping it would quietly report a group size
        # that does not deliver the stated 80% power.
        _pa = calculate_sample_size_power_analysis(effect_size=0.8, power=0.80, alpha=0.05)
        _ctrl_n = _pa.get('n_per_group', 8) if _pa.get('success') else 8
        _ctrl_range = f"{_ctrl_n}-{math.ceil(_ctrl_n * 1.1)}"
        # Add blood warnings for control group too
        warnings = []
        if blood_calc['needed'] and blood_calc['safety_color'] != 'green':
            warnings.append(f"Blood collection: {blood_calc['safety_assessment']}")

        return {
            "num_mice": group.get('num_mice', ''),
            "summary": build_summary(_ctrl_range),
            "sample_size_for_group": _ctrl_range,
            "recommended_mice": _ctrl_range,
            "toxicity_risk": 0,
            "rationale": f"Control group: recommended {_ctrl_range} animals per group (balanced with treatment groups).",
            "predicted_outcome": "No pharmacological effect expected.",
            "reference_papers": ref_corpus.get("all_papers", [])[:15],
            "validation_score": 85,
            "warnings": warnings,
            "suggested_corrections": [],
            "statistical_test": recommend_statistical_test(all_groups_count),
            "blood_calculation": blood_calc,  # ✅ NOW INCLUDED FOR CONTROL
            "ml_prediction": None,
            "all_sources": {
                "pubmed_count": len(ref_corpus.get("pubmed", [])),
                "europe_pmc_count": len(ref_corpus.get("europe_pmc", [])),
                "semantic_scholar_count": len(ref_corpus.get("semantic_scholar", [])),
                "openalex_count": len(ref_corpus.get("openalex", [])),
                "crossref_count": len(ref_corpus.get("crossref", [])),
                "doaj_count": len(ref_corpus.get("doaj", [])),
                "preprint_count": len(ref_corpus.get("preprint", [])),
                "impc_count": len(ref_corpus.get("impc", [])),
                "total_papers": len(ref_corpus.get("all_papers", []))
            }
        }

    # Sample size comes from a formal two-sample power analysis and nothing
    # else. Every assumption is stated back to the researcher so an IACUC
    # reviewer can reproduce the number independently.
    ml_prediction = None
    
    # The effect size is the whole calculation. Where the researcher names a
    # measurable endpoint and the difference worth detecting, d comes from the
    # predicted variability of that endpoint; otherwise it falls back to Cohen's
    # convention, and the basis says which of the two produced the number.
    _measure = (group.get('primary_endpoint') or '').strip()
    _target_pct = parse_float_safe(group.get('detectable_difference_pct'), 0)
    _cv = predict_outcome_cv(_measure, group.get('species'),
                             group.get('strain'), group.get('sex')) if _measure else None

    if _cv and _cv.get('in_domain') and _target_pct > 0:
        effect_size = (_target_pct / 100.0) / (_cv['cv_pct'] / 100.0)
        basis_kind = 'predicted variability'
    else:
        effect_size = 0.8
        basis_kind = 'default assumption'

    power_result = calculate_sample_size_power_analysis(
        effect_size=effect_size,
        power=0.80,
        alpha=0.05
    )
    suggested_n = power_result['n_per_group']
    sample_size_basis = {
        'method': 'Two-sample t-test power analysis (statsmodels tt_ind_solve_power)',
        'effect_size_d': round(effect_size, 3),
        'effect_size_from': basis_kind,
        'power': 0.80,
        'alpha': 0.05,
        'n_per_group': suggested_n,
        'attrition_allowance': '10%',
        'variability': _cv,
        'primary_endpoint': _measure or None,
        'detectable_difference_pct': _target_pct or None,
    }
    if basis_kind == 'predicted variability':
        sample_size_basis['note'] = (
            f"Detecting a {_target_pct:g}% difference in {_measure} at a predicted "
            f"CV of {_cv['cv_pct']}% gives d = {effect_size:.2f}.")
    elif _measure and _cv and not _cv.get('in_domain'):
        sample_size_basis['note'] = (
            f"'{_measure}' is not in the reference databases, so its variability "
            "was not predicted. N falls back to Cohen's d = 0.8 — a convention, "
            "not a measurement for this endpoint.")
    elif _measure and not _target_pct:
        sample_size_basis['note'] = (
            "State the difference worth detecting to base N on the predicted "
            "variability; until then N uses Cohen's d = 0.8.")
    else:
        sample_size_basis['note'] = (
            "No primary endpoint given, so N uses Cohen's d = 0.8 — a convention, "
            "not a measurement. Naming the endpoint and the difference worth "
            "detecting bases N on published variability instead.")
    
    # Generate power curve data
    sample_sizes, powers = generate_power_curve_data(effect_size)
    power_curve = {
        'sample_sizes': sample_sizes,
        'powers': powers,
        'effect_size': effect_size
    }
    
    n_with_attrition = math.ceil(suggested_n * 1.1)
    
    # Rationale states the method and its assumptions, nothing more.
    rationale_parts = []
    rationale_parts.append(
        f"Recommended sample size: {suggested_n}-{n_with_attrition} {animal_word} per group.")
    rationale_parts.append(
        f"Two-sample t-test power analysis: {int(0.80 * 100)}% power, alpha 0.05, "
        f"Cohen's d {effect_size}, plus a 10% attrition allowance.")
    if has_refs:
        rationale_parts.append(f"Found {len(ref_corpus['all_papers'])} relevant papers across multiple databases.")
    
    rationale = " ".join(rationale_parts)
    
    # Validation
    warnings = []
    corrections = []
    score = 70
    
    if len(ref_corpus['all_papers']) > 5:
        score += 15
    elif len(ref_corpus['all_papers']) > 0:
        score += 10
    
    num_mice = parse_int_safe(group.get('num_mice'), 0)
    if num_mice:
        diff = num_mice - suggested_n
        if -1 <= diff <= 1:
            score += 15
        elif abs(diff) > 3:
            warnings.append(f"Sample size differs significantly from recommendation (±{abs(diff)} mice).")
            corrections.append(f"Consider adjusting to {suggested_n}-{n_with_attrition} mice.")
    
    # Blood warnings
    if blood_calc['needed'] and blood_calc['safety_color'] != 'green':
        warnings.append(f"Blood collection: {blood_calc['safety_assessment']}")
    
    score = max(0, min(100, score))
    
        # Blood warnings
    if blood_calc['needed'] and blood_calc['safety_color'] != 'green':
        warnings.append(f"Blood collection: {blood_calc['safety_assessment']}")

    score = max(0, min(100, score))

    # Take the first paper from the corpus and build a safe link to it
    all_papers_list = ref_corpus.get('all_papers', [])
    first_paper = all_papers_list[0] if all_papers_list else None
    if first_paper:
        paper_url = (
            first_paper.get('url')
            or first_paper.get('pubmed_url')
            or (f"https://doi.org/{first_paper.get('doi')}" if first_paper.get('doi') else None)
            or "#"
        )
    else:
        paper_url = "#"

    return {
        "num_mice": group.get('num_mice', ''),
        "summary": build_summary(f"{suggested_n}-{n_with_attrition}"),
        "sample_size_for_group": f"{suggested_n}-{n_with_attrition} {animal_word} per group",
        "recommended_mice": f"{suggested_n}-{n_with_attrition}",
        "sample_size_details": power_result,
        "power_curve": power_curve,
        "toxicity_risk": 15,
        "rationale": rationale,
        "predicted_outcome": "Based on statistical power analysis and literature review.",
        "reference_papers": all_papers_list[:15],
        "validation_score": score,
        "warnings": warnings,
        "suggested_corrections": corrections,
        "statistical_test": recommend_statistical_test(all_groups_count),
        "blood_calculation": blood_calc,
        "ml_prediction": ml_prediction,
        "sample_size_basis": sample_size_basis,
        "all_sources": {
            "pubmed_count": len(ref_corpus.get("pubmed", [])),
            "europe_pmc_count": len(ref_corpus.get("europe_pmc", [])),
            "semantic_scholar_count": len(ref_corpus.get("semantic_scholar", [])),
            "openalex_count": len(ref_corpus.get("openalex", [])),
            "crossref_count": len(ref_corpus.get("crossref", [])),
            "doaj_count": len(ref_corpus.get("doaj", [])),
            "preprint_count": len(ref_corpus.get("preprint", [])),
            "impc_count": len(ref_corpus.get("impc", [])),
            "total_papers": len(all_papers_list)
        },
        "paper_url": paper_url,
        "source": f"{len(all_papers_list)} papers from multiple sources"
    }

# ============================================================================
# ENHANCED PDF GENERATION WITH DIAGRAMS
# ============================================================================






# ── Study flow chart (shared by the PDF and Word study-plan exports) ─────────
# One unified flow chart for the whole study, regardless of how many groups.
FLOW_PALETTE_HEX = ['#4CAF50', '#2196F3', '#FF9800', '#9C27B0', '#009688', '#E91E63', '#3F51B5']
FLOW_PALETTE_RGB = [(76, 175, 80), (33, 150, 243), (255, 152, 0), (156, 39, 176),
                    (0, 150, 136), (233, 30, 99), (63, 81, 181)]


def derive_flow_phases(study_data):
    """Ordered, de-duplicated phase names for the study flow chart.

    Prefers the phase labels from the (possibly edited) day-by-day timelines —
    unioned across every group so one flow chart covers the whole study — and
    falls back to a standard protocol sequence inferred from the group data.
    """
    seen = []
    for tl in (study_data.get('timelines') or []):
        for step in (tl.get('steps') or []):
            p = (step.get('phase') or '').strip()
            if p and p not in seen:
                seen.append(p)
    if seen:
        return seen[:7]
    # Fallback: infer a standard protocol sequence from the groups themselves
    groups = study_data.get('groups') or []
    has_drug = any((g.get('drug_name') or g.get('drug')) for g in groups)
    has_samples = any((g.get('sample_types') or g.get('toxicity_endpoints')) for g in groups)
    flow = ['Acclimatization', 'Baseline', 'Dosing' if has_drug else 'Intervention', 'Monitoring']
    if has_samples:
        flow.append('Sample collection')
    flow.append('Endpoint')
    return flow[:7]


def build_flowchart_flowables(phase_names, heading_style):
    """Return reportlab flowables: a header box + a horizontal phase flow chart
    (coloured boxes joined by arrows) that always fits the usable page width."""
    if not phase_names:
        return []
    n = len(phase_names)
    arrow_w = 0.28
    box_w = max(0.7, (7.0 - arrow_w * (n - 1)) / n)   # fit within 7" usable width
    box_style = ParagraphStyle('fcbox', fontName='Helvetica-Bold', fontSize=8,
                               leading=10, textColor=colors.white, alignment=TA_CENTER)
    arrow_style = ParagraphStyle('fcarrow', fontName='Helvetica-Bold', fontSize=13,
                                 textColor=HexColor('#26a65b'), alignment=TA_CENTER)
    row, widths, bg = [], [], []
    col = 0
    for i, name in enumerate(phase_names):
        row.append(Paragraph(name, box_style))
        widths.append(box_w * inch)
        bg.append(('BACKGROUND', (col, 0), (col, 0),
                   HexColor(FLOW_PALETTE_HEX[i % len(FLOW_PALETTE_HEX)])))
        col += 1
        if i < n - 1:
            row.append(Paragraph('→', arrow_style))
            widths.append(arrow_w * inch)
            col += 1
    chart = Table([row], colWidths=widths)
    chart.setStyle(TableStyle(bg + [
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 7),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
    ]))
    header = Table([[Paragraph("🔀 Study Flow Chart", heading_style)]], colWidths=[7 * inch])
    header.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#ede9fe')),
        ('BOX', (0, 0), (-1, -1), 1, HexColor('#8b5cf6')),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    return [header, Spacer(1, 0.05 * inch), KeepTogether(chart), Spacer(1, 0.12 * inch)]


def add_flowchart_docx(doc, phase_names):
    """Add the unified study flow chart (coloured phase boxes + arrows) to Word."""
    if not phase_names:
        return
    doc.add_heading("Study Flow Chart", level=1)
    ncols = 2 * len(phase_names) - 1
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    table.allow_autofit = False
    # Force a fixed column layout so boxes stay wide and arrows stay narrow.
    _tblPr = table._tbl.tblPr
    _layout = OxmlElement('w:tblLayout')
    _layout.set(qn('w:type'), 'fixed')
    _tblPr.append(_layout)
    # Fit within the ~6.5" usable portrait width.
    n = len(phase_names)
    box_w = Inches(min(1.1, (6.4 - 0.2 * (n - 1)) / n))
    ci = 0
    for i, name in enumerate(phase_names):
        cell = table.rows[0].cells[ci]
        cell.text = name
        try:
            cell.width = box_w
        except Exception:
            pass
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '%02x%02x%02x' % FLOW_PALETTE_RGB[i % len(FLOW_PALETTE_RGB)])
        cell._element.get_or_add_tcPr().append(shd)
        ci += 1
        if i < len(phase_names) - 1:
            acell = table.rows[0].cells[ci]
            acell.text = '→'
            try:
                acell.width = Inches(0.2)
            except Exception:
                pass
            ap = acell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in ap.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(38, 166, 91)
            ci += 1
    doc.add_paragraph()


def generate_enhanced_pdf(study_data):
    """Generate enhanced PDF with backgrounds and optimized to fit on one page."""
    buffer = BytesIO()
    # Optimize margins for better fit
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           topMargin=0.4*inch, bottomMargin=0.4*inch,
                           leftMargin=0.5*inch, rightMargin=0.5*inch)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles - more compact
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=HexColor('#26a65b'),  # Green color
        spaceAfter=3,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=8,
        textColor=HexColor('#666666'),
        alignment=TA_CENTER,
        spaceAfter=2
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=10,
        textColor=HexColor('#1f8f4e'),  # Dark green
        spaceAfter=3,
        spaceBefore=6,
        fontName='Helvetica-Bold',
        leftIndent=0
    )
    
    # Title section with background - very compact
    story.append(Spacer(1, 0.1*inch))
    
    # Add title with background box
    title_text = study_data.get('study_title', 'Untitled Study')
    title_data = [[Paragraph(f"🐭 {title_text}", title_style)]]
    title_table = Table(title_data, colWidths=[7*inch])
    title_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#e9f8ef')),  # Light green background
        ('BOX', (0, 0), (-1, -1), 2, HexColor('#26a65b')),  # Green border
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
    ]))
    story.append(title_table)
    story.append(Spacer(1, 0.05*inch))
    
    # Generated date only (removed "Study Plan Document")
    subtitle_data = [[Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", subtitle_style)]]
    subtitle_table = Table(subtitle_data, colWidths=[7*inch])
    subtitle_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f3f7fa')),  # Light background
        ('BOX', (0, 0), (-1, -1), 1, HexColor('#cfe9dc')),  # Light border
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ]))
    story.append(subtitle_table)
    story.append(Spacer(1, 0.12*inch))
    
    groups = study_data.get('groups', [])
    
    # Determine start date from groups or use provided date
    start_date_str = study_data.get('study_start_date')
    if not start_date_str or start_date_str == '':
        group_dates = [g.get('start_date') for g in groups if g.get('start_date')]
        if group_dates:
            start_date_str = min(group_dates)
    
    phases = calculate_study_phases(groups, start_date_str)
    
    # Study Information Section with section header in colored box
    header_data = [[Paragraph("📋 Study Information", heading_style)]]
    header_table = Table(header_data, colWidths=[7*inch])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#d4edda')),  # Light green background
        ('BOX', (0, 0), (-1, -1), 1, HexColor('#26a65b')),  # Green border
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 0.05*inch))
    
    # Calculate total weeks dynamically from phases using the actual calculated durations
    total_weeks = sum([p.get('duration_weeks', 6) for p in phases])
    
    # Get start date
    display_start_date = ''
    if start_date_str:
        try:
            if isinstance(start_date_str, str):
                parsed_date = datetime.strptime(start_date_str, '%Y-%m-%d')
                display_start_date = parsed_date.strftime('%b %d, %Y')
        except:
            display_start_date = ''
    
    if not display_start_date and phases:
        display_start_date = phases[0]['start_date']
    
    info_data = [
        ['Principal Investigator:', study_data.get('pi_name', 'Not specified')],
        ['Institution:', study_data.get('institution', 'Not specified')],
        ['Study Type:', 'Rodent Pharmacology Study'],
        ['Start Date:', display_start_date if display_start_date else 'TBD']
    ]
    
    # More compact table with enhanced background
    info_table = Table(info_data, colWidths=[2*inch, 5*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#e9f8ef')),  # Light green background
        ('TEXTCOLOR', (0, 0), (0, -1), HexColor('#155724')),  # Dark green labels
        ('TEXTCOLOR', (1, 0), (1, -1), colors.black),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#cfe9dc')),
        ('BOX', (0, 0), (-1, -1), 2, HexColor('#26a65b'))
    ]))
    
    story.append(info_table)
    story.append(Spacer(1, 0.1*inch))

    # Unified study flow chart (all groups, one chart)
    story.extend(build_flowchart_flowables(derive_flow_phases(study_data), heading_style))

    # Editable day-by-day protocol timeline(s)
    # ONE unified timeline for the whole study (all groups merged).
    timelines = study_data.get('timelines', []) or []
    merged_steps = merge_timelines(timelines)
    if merged_steps:
        tl_header = [[Paragraph("🗓️ Experimental Timeline / Study Plan (all groups)", heading_style)]]
        tl_header_table = Table(tl_header, colWidths=[7*inch])
        tl_header_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#dbeafe')),
            ('BOX', (0, 0), (-1, -1), 1, HexColor('#2563eb')),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(tl_header_table)
        story.append(Spacer(1, 0.04*inch))

        cell_style = ParagraphStyle('tlcell', fontName='Helvetica', fontSize=8, leading=10)
        head_style = ParagraphStyle('tlhead', fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.white)
        tl_data = [[Paragraph('Day', head_style), Paragraph('Phase', head_style), Paragraph('Activity', head_style)]]
        for step in merged_steps:
            # per-group lines are newline-separated -> <br/> for reportlab
            activity = str(escape(step.get('activity', ''))).replace('\n', '<br/>')
            tl_data.append([
                Paragraph(str(escape(step.get('day', ''))), cell_style),
                Paragraph(str(escape(step.get('phase', ''))), cell_style),
                Paragraph(activity, cell_style),
            ])
        tl_table = Table(tl_data, colWidths=[0.9*inch, 1.3*inch, 4.8*inch])
        tl_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#2563eb')),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#bfdbfe')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(tl_table)
        story.append(Spacer(1, 0.12*inch))

    # Special instructions (per group) — shown prominently right after timeline
    special_rows = [(g.get('group_name', f'Group {i}'), (g.get('instructions') or '').strip())
                    for i, g in enumerate(groups, 1)]
    special_rows = [(n, ins) for n, ins in special_rows if ins]
    if special_rows:
        si_header = [[Paragraph("📌 Special Instructions", heading_style)]]
        si_ht = Table(si_header, colWidths=[7*inch])
        si_ht.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#fffbeb')),
            ('BOX', (0, 0), (-1, -1), 1, HexColor('#f59e0b')),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(si_ht)
        story.append(Spacer(1, 0.04*inch))
        si_cell = ParagraphStyle('si', fontName='Helvetica', fontSize=8, leading=11)
        si_lbl = ParagraphStyle('sil', fontName='Helvetica-Bold', fontSize=8, leading=11)
        si_data = [[Paragraph(escape(n), si_lbl), Paragraph(escape(ins), si_cell)]
                   for n, ins in special_rows]
        si_table = Table(si_data, colWidths=[1.8*inch, 5.2*inch])
        si_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#fffef5')),
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#fde68a')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6), ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(si_table)
        story.append(Spacer(1, 0.12*inch))

    # Experimental Groups - Detailed Analysis Section
    if groups:
        groups_header_data = [[Paragraph("🧪 Experimental Groups - Detailed Analysis", heading_style)]]
        groups_header_table = Table(groups_header_data, colWidths=[7*inch])
        groups_header_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor('#d4edda')),
            ('BOX', (0, 0), (-1, -1), 1, HexColor('#26a65b')),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(groups_header_table)
        story.append(Spacer(1, 0.05*inch))
        
        # Add each group as a detailed table
        for idx, g in enumerate(groups, 1):
            # Group title
            group_title = Paragraph(f"<b>Group {idx}: {g.get('group_name', f'Group {idx}')}</b>",
                                   ParagraphStyle('GroupTitle', parent=styles['Normal'], fontSize=10,
                                                fontName='Helvetica-Bold', textColor=HexColor('#1f8f4e')))
            story.append(group_title)
            story.append(Spacer(1, 0.05*inch))
            
            # Group details table
            group_details = [
                ['Drug Name:', g.get('drug_name', g.get('drug', 'N/A'))],
                ['Dose:', f"{g.get('dose', 'N/A')} mg/kg"],
                ['Strain:', g.get('strain', 'N/A')],
                ['Number of Mice:', str(g.get('num_mice', 'N/A'))],
                ['Route:', g.get('route', 'N/A')],
                ['Target Organ:', g.get('target_organ', 'N/A')],
                ['Start Date:', g.get('start_date', 'TBD')]
            ]
            
            group_table = Table(group_details, colWidths=[2*inch, 5*inch])
            group_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f9fafb')),
                ('TEXTCOLOR', (0, 0), (0, -1), HexColor('#374151')),
                ('TEXTCOLOR', (1, 0), (1, -1), colors.black),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#e5e7eb')),
                ('BOX', (0, 0), (-1, -1), 1, HexColor('#9ca3af'))
            ]))
            story.append(group_table)
            
            # Add special instructions if present
            if g.get('instructions'):
                story.append(Spacer(1, 0.05*inch))
                instructions_data = [[Paragraph(f"<b>Special Instructions:</b> {g.get('instructions', '')}",
                                               ParagraphStyle('Instructions', parent=styles['Normal'], fontSize=8))]]
                instructions_table = Table(instructions_data, colWidths=[7*inch])
                instructions_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), HexColor('#fffbeb')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('BOX', (0, 0), (-1, -1), 1, HexColor('#fbbf24'))
                ]))
                story.append(instructions_table)
            
            # Add spacing between groups
            if idx < len(groups):
                story.append(Spacer(1, 0.1*inch))
        
        story.append(Spacer(1, 0.1*inch))
    
    # Add disclaimer at bottom with background
    disclaimer_data = [[Paragraph(
        "⚠️ DISCLAIMER: This is a prototype demonstration tool. Not for actual clinical use. All animal research must comply with institutional guidelines and regulations.",
        ParagraphStyle('Disclaimer', parent=styles['Normal'], fontSize=6,
                      textColor=HexColor('#666666'), alignment=TA_CENTER)
    )]]
    disclaimer_table = Table(disclaimer_data, colWidths=[7*inch])
    disclaimer_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), HexColor('#f3f4f6')),
        ('BOX', (0, 0), (-1, -1), 1, HexColor('#d1d5db')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(disclaimer_table)
    
    # Build PDF - keep everything on one page
    doc.build(story)
    buffer.seek(0)
    return buffer


def merge_timelines(timelines):
    """Merge per-group timelines into ONE unified timeline.

    Rows are keyed by (day, phase) in first-seen order. When the activity is the
    same for every group it is shown once; when it differs (e.g. the dosing step)
    each group's instruction is listed within that row so a single timeline
    carries the whole study."""
    order, acts = [], {}
    for tl in (timelines or []):
        grp = (tl.get('group') or '').strip()
        for step in (tl.get('steps') or []):
            key = (str(step.get('day', '')).strip(), str(step.get('phase', '')).strip())
            act = str(step.get('activity', '')).strip()
            if key not in acts:
                acts[key] = []
                order.append(key)
            acts[key].append((grp, act))
    merged = []
    for key in order:
        pairs = acts[key]
        distinct = list(dict.fromkeys(a for _, a in pairs))
        if len(distinct) <= 1:
            activity = distinct[0] if distinct else ''
        else:
            lines = [f"{g}: {a}" if g else a for g, a in pairs]
            activity = "\n".join(dict.fromkeys(lines))   # de-dup identical lines
        merged.append({'day': key[0], 'phase': key[1], 'activity': activity})
    return merged


def add_protocol_timelines_to_docx(doc, timelines):
    """Render ONE unified day-by-day protocol timeline (all groups) into Word."""
    steps = merge_timelines(timelines)
    if not steps:
        return
    doc.add_heading("Experimental Timeline / Study Plan", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Day', 'Phase', 'Activity'
    for c in hdr:
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True
    for step in steps:
        row = table.add_row().cells
        row[0].text = str(step.get('day', ''))
        row[1].text = str(step.get('phase', ''))
        # activity may hold per-group lines separated by \n -> real line breaks
        lines = str(step.get('activity', '')).split('\n')
        para = row[2].paragraphs[0]
        para.add_run(lines[0])
        for extra in lines[1:]:
            run = para.add_run()
            run.add_break()
            para.add_run(extra)
    doc.add_paragraph()



def calculate_study_phases(groups, study_start_date=None):
    """Calculate study phases - ONE PHASE PER GROUP directly linked to group data."""
    if not groups:
        return []
    
    phases = []
    
    # Create one phase for each group
    for idx, group in enumerate(groups, 1):
        group_name = group.get('group_name', f'Group {idx}')
        drug_name = group.get('drug_name', 'Unknown')
        instructions = group.get('instructions', '').strip()
        
        # Get group start and end dates from the group data
        group_start_date = group.get('start_date', '')
        group_end_date = group.get('end_date', '')
        
        # Calculate duration and format dates
        if group_start_date and group_end_date:
            try:
                start_dt = datetime.strptime(group_start_date, '%Y-%m-%d')
                end_dt = datetime.strptime(group_end_date, '%Y-%m-%d')
                
                # Calculate actual duration in weeks
                duration_days = (end_dt - start_dt).days
                duration_weeks = max(1, round(duration_days / 7))  # At least 1 week
                
                start_formatted = start_dt.strftime('%b %d, %Y')
                end_formatted = end_dt.strftime('%b %d, %Y')
                
            except Exception as e:
                logger.warning(f"Date parsing error for group {idx}: {e}")
                start_formatted = 'TBD'
                end_formatted = 'TBD'
                duration_weeks = 6
        elif group_start_date:
            # Only start date provided, use default 6 weeks
            try:
                start_dt = datetime.strptime(group_start_date, '%Y-%m-%d')
                start_formatted = start_dt.strftime('%b %d, %Y')
                
                duration_weeks = 6
                end_dt = start_dt + timedelta(weeks=duration_weeks)
                end_formatted = end_dt.strftime('%b %d, %Y')
                
            except Exception as e:
                logger.warning(f"Date parsing error for group {idx}: {e}")
                start_formatted = 'TBD'
                end_formatted = 'TBD'
                duration_weeks = 6
        else:
            # No dates provided
            start_formatted = 'TBD'
            end_formatted = 'TBD'
            duration_weeks = 6
        
        # Create phase directly from group
        phase = {
            'name': f'Phase {idx}',
            'title': f'{group_name} - {drug_name}',
            'duration': f'{duration_weeks} weeks',
            'duration_weeks': duration_weeks,  # Store numeric value for calculations
            'start_date': start_formatted,
            'end_date': end_formatted,
            'description': instructions,
            'group_id': idx,
            'dose': group.get('dose', ''),
            'route': group.get('route', ''),
            'num_mice': group.get('num_mice', '')
        }
        
        phases.append(phase)
    
    return phases

def generate_enhanced_docx(study_data):
    """Generate enhanced Word document matching the study plan format."""
    doc = Document()
    
    # Title
    title = doc.add_heading(study_data.get('study_title', 'Untitled Study'), level=0)
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle_para = doc.add_paragraph("Study Plan Document")
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.runs[0]
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Generated date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.runs[0]
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # Calculate study phases
    groups = study_data.get('groups', [])
    start_date = study_data.get('study_start_date', None)
    phases = calculate_study_phases(groups, start_date)
    
    # Calculate total weeks dynamically from phases using the actual calculated durations
    total_weeks = sum([p.get('duration_weeks', 6) for p in phases])
    
    # Study Information Section (with green background)
    doc.add_heading("Study Information", level=1)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid'

    info_data = [
        ['Principal Investigator:', study_data.get('pi_name', 'Not specified')],
        ['Institution:', study_data.get('institution', 'Not specified')],
        ['Study Type:', 'Rodent Pharmacology Study'],
        ['Start Date:', phases[0]['start_date'] if phases else 'TBD']
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value

        # Make labels bold
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        # Add green background to all cells
        for cell in info_table.rows[i].cells:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'd4edda')  # Light green
            cell._element.get_or_add_tcPr().append(shading_elm)

    # Unified study flow chart (all groups, one chart)
    doc.add_paragraph()
    add_flowchart_docx(doc, derive_flow_phases(study_data))

    # Editable day-by-day protocol timeline (from the results / user edits)
    doc.add_paragraph()
    add_protocol_timelines_to_docx(doc, study_data.get('timelines', []))

    # Special instructions (per group) — prominent, right after the timeline
    special_rows = [(g.get('group_name', f'Group {i}'), (g.get('instructions') or '').strip())
                    for i, g in enumerate(groups, 1)]
    special_rows = [(n, ins) for n, ins in special_rows if ins]
    if special_rows:
        doc.add_paragraph()
        doc.add_heading("📌 Special Instructions", level=1)
        si_tbl = doc.add_table(rows=len(special_rows), cols=2)
        si_tbl.style = 'Light Grid Accent 1'
        for i, (n, ins) in enumerate(special_rows):
            si_tbl.rows[i].cells[0].text = n
            if si_tbl.rows[i].cells[0].paragraphs[0].runs:
                si_tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            si_tbl.rows[i].cells[1].text = ins
            for cell in si_tbl.rows[i].cells:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'fffbeb')
                cell._element.get_or_add_tcPr().append(shd)

    doc.add_paragraph()

    doc.add_paragraph()
    
    # Disclaimer
    disclaimer_para = doc.add_paragraph(
        "⚠️ Research Protocol: This document is for planning purposes only. "
        "Always consult with qualified researchers and obtain proper ethical approvals."
    )
    disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer_run = disclaimer_para.runs[0]
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Page 2: Detailed Group Information
    if groups:
        doc.add_page_break()
        doc.add_heading("🧪 Experimental Groups - Detailed Analysis", level=1)
        
        for idx, g in enumerate(groups, 1):
            doc.add_heading(f"Group {idx}: {g.get('group_name', '')}", level=2)
            
            # Group details table
            group_table = doc.add_table(rows=8, cols=2)
            group_table.style = 'Light List Accent 1'
            
            details = [
                ["Drug Name", g.get('drug_name', g.get('drug', ''))],
                ["Dose", f"{g.get('dose', '')} mg/kg"],
                ["Strain", g.get('strain', '')],
                ["Number of Mice", g.get('num_mice', '')],
                ["Route", g.get('route', '')],
                ["Target Organ", g.get('target_organ', '')],
                ["Start Date", g.get('start_date', 'TBD')]
            ]
            
            for i, (label, value) in enumerate(details):
                group_table.rows[i].cells[0].text = label
                group_table.rows[i].cells[1].text = str(value)
                group_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Add instructions in merged row if present
            if g.get('instructions'):
                group_table.rows[7].cells[0].text = "Special Instructions"
                group_table.rows[7].cells[0].paragraphs[0].runs[0].font.bold = True
                group_table.rows[7].cells[1].text = g.get('instructions', '')
                
                # Add light yellow background to instructions
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'fffbeb')
                group_table.rows[7].cells[1]._element.get_or_add_tcPr().append(shading_elm)
            
            doc.add_paragraph()
            
            # Blood calculation if needed
            if 'blood_calculation' in g and g['blood_calculation'].get('needed'):
                doc.add_heading("💉 Blood Collection Requirements", level=3)
                
                blood_info = g['blood_calculation']
                blood_para = doc.add_paragraph()
                blood_para.add_run(f"Total Volume Needed: ").bold = True
                blood_para.add_run(f"{blood_info['total_volume_ml']} mL\n")
                blood_para.add_run(f"Safe Volume: ").bold = True
                blood_para.add_run(f"{blood_info['safe_volume_ml']} mL\n")
                blood_para.add_run(f"Safety: ").bold = True
                blood_para.add_run(blood_info['safety_assessment'])
                
                # Add color coding
                safety_color = blood_info.get('safety_color', 'grey')
                if safety_color == 'red':
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'fee2e2')
                    blood_para._element.get_or_add_pPr().append(shading_elm)
                elif safety_color == 'orange':
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'fef3c7')
                    blood_para._element.get_or_add_pPr().append(shading_elm)
                elif safety_color == 'green':
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'd1fae5')
                    blood_para._element.get_or_add_pPr().append(shading_elm)
                
                doc.add_paragraph()
        
        # Literature sources
        if 'all_sources' in g:
            sources = g['all_sources']
            doc.add_heading("📚 Literature Sources", level=3)
            
            sources_table = doc.add_table(rows=7, cols=2)
            sources_table.style = 'Light Grid Accent 1'
            
            source_data = [
                ["Database", "Papers Found"],
                ["PubMed", str(sources.get('pubmed_count', 0))],
                ["Semantic Scholar", str(sources.get('semantic_scholar_count', 0))],
                ["OpenAlex", str(sources.get('openalex_count', 0))],
                ["CrossRef", str(sources.get('crossref_count', 0))],
                ["Europe PMC", str(sources.get('europe_pmc_count', 0))],
                ["IMPC", str(sources.get('impc_count', 0))]
            ]
            
            for i, (label, value) in enumerate(source_data):
                sources_table.rows[i].cells[0].text = label
                sources_table.rows[i].cells[1].text = value
                if i == 0:
                    sources_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                    sources_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
        
        if idx < len(groups):
            doc.add_paragraph()
    
    # Save
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================================================
# FLASK ROUTES
# ============================================================================

@app.route('/')
def home():
    """Serve main application page."""
    return render_template('index.html')

@app.route('/variability-measures', methods=['GET'])
def variability_measures():
    """Outcome measures the variability model actually covers.

    Only these can drive the sample size; anything else falls back to the
    conventional effect size, so the form offers them explicitly rather than
    letting a researcher type something the model will silently decline.
    """
    _load_ovm()
    items = _OVM_MEASURES or []
    species = (request.args.get('species') or '').strip().lower()
    if species:
        items = [m for m in items if m['species'] == species]
    return jsonify({'count': len(items), 'measures': items[:400]})


@app.route('/sample-types', methods=['GET'])
def get_sample_types():
    """Return list of available sample types."""
    return jsonify({"sample_types": SAMPLE_TYPES})

@app.route('/apply-samples-to-all', methods=['POST'])
def apply_samples_to_all():
    """Apply selected samples to all groups."""
    try:
        data = request.get_json()
        sample_types = data.get('sample_types', [])
        groups = data.get('groups', [])
        
        # Apply sample types to all groups
        for group in groups:
            group['sample_types'] = sample_types.copy()
        
        return jsonify({
            "success": True,
            "message": f"Applied {len(sample_types)} sample type(s) to {len(groups)} group(s)",
            "groups": groups
        })
    
    except Exception as e:
        logger.exception("Error applying samples to all groups")
        return jsonify({"error": str(e)}), 500

@app.route('/predict', methods=['POST'])
@limiter.limit(Config.RATE_LIMIT_PREDICT) if limiter else lambda f: f
def predict():
    """Main prediction endpoint with enhanced features."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
        
        if VALIDATION_AVAILABLE:
            try:
                validated_data = StudySchema().load(data)
                all_groups = validated_data.get('groups', [])
            except ValidationError as err:
                return jsonify({"error": "Validation failed", "details": err.messages}), 400
        else:
            all_groups = data.get('groups', [])
        
        if not all_groups:
            return jsonify({"error": "No groups provided"}), 400
        
        results = []
        all_groups_count = len(all_groups)
        
        for group in all_groups:
            gname = group.get('group_name', 'Group')
            try:
                if is_confidential(group):
                    # Never resolve a confidential name against PubChem. The
                    # investigator supplies the structure instead.
                    if not (group.get('smiles') or '').strip():
                        results.append({"group_name": gname,
                                        "error": "Confidential compound: please paste the chemical "
                                                 "structure (SMILES) so the predictions can run locally."})
                        continue
                    ncbi = {"success": True, "ncbi_link": "#"}
                else:
                    ncbi = get_drug_data_from_ncbi(group.get('drug_name', ''))

                if not ncbi['success'] and not ncbi.get('is_control'):
                    results.append({"group_name": gname, "error": ncbi['error']})
                    continue

                suggestion = get_prediction_and_suggestion(group, all_groups_count)

                results.append({
                    "group_name": gname,
                    "drug": (group.get('drug_name') or '').capitalize(),
                    "ncbi_link": ncbi.get('ncbi_link', '#'),
                    **suggestion
                })
            except Exception as ge:
                # One group failing must not break the whole analysis.
                logger.exception(f"Error analyzing group '{gname}'")
                results.append({
                    "group_name": gname,
                    "error": f"Could not analyze this group ({type(ge).__name__}). "
                             f"Try again, or check the drug name and dose."
                })

        return jsonify(results)
    
    except Exception as e:
        logger.exception("Error in predict endpoint")
        return jsonify({"error": "Internal server error", "details": str(e)}), 500

@app.route('/study-plan/pdf', methods=['POST'])
def study_plan_pdf():
    """Generate enhanced PDF with diagrams."""
    try:
        data = request.get_json() or {}
        buffer = generate_enhanced_pdf(data)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name="enhanced_study_plan.pdf",
            mimetype="application/pdf"
        )
    
    except Exception as e:
        logger.exception("Error generating PDF")
        return jsonify({"error": "PDF generation failed", "details": str(e)}), 500


@app.route('/study-plan/docx', methods=['POST'])
def study_plan_docx():
    """Generate enhanced Word document."""
    try:
        data = request.get_json() or {}
        buffer = generate_enhanced_docx(data)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name="enhanced_study_plan.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    except Exception as e:
        logger.exception("Error generating DOCX")
        return jsonify({"error": "DOCX generation failed", "details": str(e)}), 500


@app.route('/iacuc/generate', methods=['POST'])
def iacuc_generate():
    """Auto-fill the KAIMRC IACUC Animal Ethics form (.docx) from the analysis.

    Body (JSON): {study, groups, analysis, admin} — see iacuc_generator.
    The scientific narrative is synthesised from the platform's analysis; the
    researcher supplies only administrative details (team, funding, housing).
    Nothing is persisted.
    """
    try:
        payload = request.get_json(silent=True) or {}
        if not payload.get("groups"):
            return jsonify({"error": "No study groups provided. Run an analysis first."}), 400
        buffer = generate_iacuc_docx(payload)
        title = (payload.get("study") or {}).get("study_title") or "study"
        safe = re.sub(r'[^A-Za-z0-9_-]+', '_', title).strip('_')[:60] or "study"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"IACUC_{safe}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except FileNotFoundError as e:
        logger.error(f"IACUC template missing: {e}")
        return jsonify({"error": "IACUC template not found on server."}), 500
    except Exception as e:
        logger.exception("Error generating IACUC form")
        return jsonify({"error": "IACUC generation failed", "details": str(e)}), 500


@app.route('/calculate-blood', methods=['POST'])
def calculate_blood():
    """Calculate blood quantity requirements."""
    try:
        data = request.get_json()
        
        result = BloodQuantityCalculator.calculate_blood_needed(
            sample_types=data.get('sample_types', []),
            weight_g=data.get('weight', 25),
            timepoints=data.get('timepoints', 1),
            num_replicates=data.get('replicates', 1)
        )
        
        return jsonify(result)
    
    except Exception as e:
        logger.exception("Error calculating blood quantity")
        return jsonify({"error": str(e)}), 500

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Endpoint not found"}), 404

@app.errorhandler(500)
def internal_error(error):
    logger.exception("Internal server error")
    return jsonify({"error": "Internal server error"}), 500

@app.errorhandler(429)
def ratelimit_handler(e):
    return jsonify({"error": "Rate limit exceeded"}), 429









# ============================================================================
# NEW HELPER FUNCTIONS FOR COMPREHENSIVE DRUG ANALYSIS
# ============================================================================



# GHS acute oral toxicity classification (UN GHS Rev.10, Table 3.1.1) —
# the cut-off values are the published regulatory bands, not a scale of ours.
# (upper LD50 bound mg/kg, category label, GHS hazard statement, internal risk score)
# upper bound, GHS category, dose band, the category's own hazard word, H code,
# internal risk weight used only to scale welfare monitoring.
_GHS_ORAL_BANDS = [
    (5,     'Category 1', '≤ 5 mg/kg',         'Fatal',         'H300', 90),
    (50,    'Category 2', '> 5–50 mg/kg',      'Fatal',         'H300', 80),
    (300,   'Category 3', '> 50–300 mg/kg',    'Toxic',         'H301', 65),
    (2000,  'Category 4', '> 300–2000 mg/kg',  'Harmful',       'H302', 45),
    (5000,  'Category 5', '> 2000–5000 mg/kg', 'May be harmful', 'H303', 30),
]


# The published source behind every acute-toxicity label the platform shows, so
# a reviewer can check the band rather than take the classification on trust.
_GHS_REFERENCE = {
    'citation': ('United Nations. Globally Harmonized System of Classification and '
                 'Labelling of Chemicals (GHS), Rev.10 (2023), Chapter 3.1 "Acute '
                 'toxicity", Table 3.1.1 — acute toxicity hazard categories and '
                 'the (approximate) LD50 values defining them.'),
    'url': 'https://unece.org/transport/dangerous-goods/ghs-rev10-2023',
    'pdf_en': 'https://unece.org/sites/default/files/2023-07/GHS%20Rev10e.pdf',
    'publisher': 'UNECE',
}


def ghs_scale_table():
    """Every oral band, so the researcher sees where their compound falls."""
    rows = [{'category': label, 'band': band, 'word': word, 'code': code}
            for _, label, band, word, code, _ in _GHS_ORAL_BANDS]
    rows.append({'category': 'Not classified', 'band': '> 5000 mg/kg',
                 'word': 'Not classified', 'code': '—'})
    return rows


def ld50_to_ghs(ld50_mg_kg):
    """Classify an oral LD50 into its GHS acute toxicity category.

    Returns the published regulatory band — no invented percentage. `risk_score`
    is an INTERNAL value used only to scale welfare-monitoring intensity; it is
    never displayed as a measurement.
    """
    try:
        ld50 = float(ld50_mg_kg)
    except (TypeError, ValueError):
        return None
    for upper, label, band, word, code, risk in _GHS_ORAL_BANDS:
        if ld50 <= upper:
            return {'category': label, 'band': band, 'word': word,
                    'hazard': f'{code} — {word} if swallowed',
                    'risk_score': risk, 'scale': 'GHS acute oral toxicity (UN GHS Rev.10)',
                    'reference': _GHS_REFERENCE, 'scale_table': ghs_scale_table()}
    return {'category': 'Not classified', 'band': '> 5000 mg/kg',
            'word': 'Not classified', 'hazard': 'Below the GHS classification threshold',
            'risk_score': 20, 'scale': 'GHS acute oral toxicity (UN GHS Rev.10)',
            'reference': _GHS_REFERENCE, 'scale_table': ghs_scale_table()}


def round_sig(x, sig=2):
    """Round to `sig` significant figures — model estimates carry a wide error
    band, so reporting extra decimals would imply precision that is not there."""
    import math
    try:
        x = float(x)
    except (TypeError, ValueError):
        return None
    if x == 0:
        return 0.0
    r = round(x, -int(math.floor(math.log10(abs(x)))) + (sig - 1))
    return int(r) if r >= 10 else r


def get_real_ld50_info(drug_name, cid=None, route='oral', species='Mouse'):
    """
    Shared helper: return the best real LD50 estimate for a drug using the
    hybrid predictor (experimental data first, trained ML model second).
    Prefers experimental values for the selected species and route.
    Returns dict with ld50_mg_kg, source, category (or found=False).
    """
    try:
        exp = toxicity_predictor.get_experimental_ld50(drug_name, cid=cid, route=route, species=species)
        if exp.get('found'):
            return {
                'found': True,
                'ld50_mg_kg': exp['ld50_mg_kg'],
                'category': exp['category'],
                'source': 'experimental',
                'source_detail': exp.get('source_detail', ''),
            }
        structure = toxicity_predictor.get_chemical_structure(drug_name)
        if structure.get('success'):
            ml = toxicity_predictor.predict_ld50_ml(structure['smiles'])
            if ml:
                return {
                    'found': True,
                    'ld50_mg_kg': ml['ld50_mg_kg'],
                    'category': ml['category'],
                    'source': 'ml_model',
                    'source_detail': ml.get('source_detail', ''),
                }
    except Exception as e:
        logger.error(f"get_real_ld50_info failed for {drug_name}: {e}")
    return {'found': False}


























@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint to verify app is running."""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '3.1 Enhanced',
        'features': {
            'ml_available': ML_AVAILABLE,
            'rdkit_available': RDKIT_AVAILABLE,
            'validation_available': VALIDATION_AVAILABLE,
            'comprehensive_drug_analysis': True
        },
        'endpoints': {
            'home': '/',
            'predict_toxicity': '/predict-toxicity',
            'predict_effectiveness': '/predict-effectiveness',
            'predict_complete': '/predict-complete',
            'analyze_drug_comprehensive': '/analyze-drug-comprehensive (NEW)',
            'test_assessment': '/test-assessment',
            'health': '/health'
        }
    })


# ============================================================================
# MAIN
# ============================================================================

if __name__ == '__main__':
    logger.info("=" * 70)
    logger.info("Starting Enhanced Rodent Study Planner v3.1")
    logger.info("=" * 70)
    logger.info(f"✓ ML features: {'ENABLED' if ML_AVAILABLE else 'DISABLED'}")
    logger.info(f"✓ Validation: {'ENABLED' if VALIDATION_AVAILABLE else 'DISABLED'}")
    logger.info(f"✓ APIs configured: PubMed, Semantic Scholar, OpenAlex, CrossRef")
    logger.info(f"✓ Blood calculator: ACTIVE")
    logger.info(f"✓ Enhanced PDF/DOCX: ACTIVE (with diagrams)")
    logger.info(f"✓ NEW: Comprehensive Drug Analysis: ACTIVE")
    logger.info(f"  - Toxicity percentage (0-100%) with PubChem")
    logger.info(f"  - Efficacy percentage (0-100%) with literature")
    logger.info(f"  - Mouse-specific modifiers (8 parameters)")
    logger.info(f"  - Risk-benefit assessment")
    logger.info(f"  - Auto-generated recommendations")
    logger.info(f"✓ NEW ENDPOINT: /analyze-drug-comprehensive")
    logger.info("=" * 70)
    logger.info("🚀 Server starting on http://0.0.0.0:5000")
    logger.info("=" * 70)
    
    app.run(
        debug=Config.DEBUG,
        host='0.0.0.0',
        port=5000
    )
