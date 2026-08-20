"""
Build the docxtpl IACUC template from the blank KAIMRC Animal Ethics form.

Run ONCE (or whenever the blank form is updated) to produce
`iacuc_assets/iacuc_template.docx`, which the Flask app renders at runtime.

It inserts Jinja tags into the exact answer cells identified by table index,
touching ONLY the text of answer paragraphs — never styles, fonts, colours,
column widths or page layout — so the rendered document is visually identical
to the original blank form: same layout, colours and fonts, delivered as Word.

Usage:
    python iacuc_assets/build_template.py "/path/to/blank_form.docx"
"""
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
DEFAULT_SRC = Path.home() / "Downloads" / \
    "Appendix A - Animal Ethics Application Form17-7-2025.docx"
OUT = HERE / "iacuc_template.docx"

# ── 1x1 free-text answer boxes: table index -> Jinja tag ──────────────────
# (The scientific narrative is generated server-side from the study data.)
BOX_TAGS = {
    6:  "study_purpose",            # PART 2 – purpose of the study
    7:  "what_done",                # PART 2 – what will be done to the animals
    8:  "expected_outcomes",        # PART 2 – expected outcomes
    9:  "study_benefit",            # PART 2 – benefit to science/medicine/society
    12: "strain_health_issues",     # PART 3 – known health/husbandry issues
    17: "scientific_justification", # PART 4 – justify use of live animals
    18: "literature_background",    # PART 4 – literature search & background
    19: "research_aims",            # PART 4 – aims of the project
    33: "procedures_overview",      # PART 6 – overview of procedures
    57: "experimental_endpoints",   # PART 8 – experimental endpoints
    58: "humane_endpoints",         # PART 8 – humane endpoints
    59: "observation_frequency",    # PART 8 – frequency of observation
    61: "endpoint_response",        # PART 8 – actions when endpoints reached
    63: "euthanasia_method",        # PART 8 – method of euthanasia
    80: "pain_monitoring_frequency",# PART 11 – monitoring frequency
    81: "discomfort_measures",      # PART 11 – measures to minimise discomfort
}

# ── Repeating tables ──────────────────────────────────────────────────────
# The form ships 6 blank data rows in each of these tables. Rather than a
# docxtpl {%tr%} loop (which the form's content-control cells break), we fill
# fixed indexed slots with conditional {{ }} output fields — plain and robust.
# Each: table index -> (iterable_name, first_data_row, n_rows, [field per col]).
ROW_SLOTS = {
    2:  ("team", 1, 6, ["name", "role", "qualifications",
                        "institution", "email", "mobile"]),     # PART 1 personnel
    10: ("animals", 1, 6, ["species", "strain", "sex",
                           "age", "total", "source"]),          # PART 3 animal table
}


def _set_para_text(p, text):
    """Set a <w:p> element's text to `text`, preserving the first run's rPr."""
    runs = p.findall(qn("w:r"))
    if runs:
        # keep first run (and its formatting), blank the rest
        ts = runs[0].findall(qn("w:t"))
        if ts:
            ts[0].text = text
            for extra in ts[1:]:
                extra.text = ""
        else:
            t = runs[0].makeelement(qn("w:t"), {})
            t.set(qn("xml:space"), "preserve")
            t.text = text
            runs[0].append(t)
        for r in runs[1:]:
            r.getparent().remove(r)
    else:
        r = p.makeelement(qn("w:r"), {})
        t = p.makeelement(qn("w:t"), {})
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)


def _left_align(p):
    """Force left alignment on a <w:p> so multi-line answers don't get the
    box's default 'justify' (which stretches words across the line)."""
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = p.makeelement(qn("w:pPr"), {})
        p.insert(0, pPr)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = pPr.makeelement(qn("w:jc"), {})
        pPr.append(jc)
    jc.set(qn("w:val"), "left")


def set_box_tag(table, text):
    """Set the text of the first cell of a box table.

    Handles answer boxes wrapped in a content control (<w:sdt>), which
    python-docx's `.cells` does not traverse — we go straight to the first
    <w:tc>/<w:p> in the table's XML.
    """
    tc = next(table._tbl.iter(qn("w:tc")))
    p = tc.find(qn("w:p"))
    if p is None:
        p = tc.makeelement(qn("w:p"), {})
        tc.append(p)
    _set_para_text(p, text)
    _left_align(p)


def set_cell_tag(cell, text):
    """Replace a normal cell's text with `text`, keeping paragraph formatting."""
    _set_para_text(cell.paragraphs[0]._p, text)


def build(src: Path):
    doc = Document(str(src))
    tables = doc.tables

    # 1) free-text answer boxes (may be wrapped in content controls)
    for idx, tag in BOX_TAGS.items():
        set_box_tag(tables[idx], "{{ %s }}" % tag)

    # 2) The personnel (T2) and animal (T10) tables have per-cell content
    #    controls (dropdowns / "click here" text). docxtpl tags render *beside*
    #    those placeholders, which looks column-shifted — so these are filled in
    #    the runtime post-process pass instead (see iacuc_generator).

    # 3) funding source — appended next to the "Other" funding cell (PART 1)
    funding_tbl = tables[5]
    for cell in funding_tbl.rows[1].cells:
        if cell.text.strip() == "Other":
            cell.paragraphs[0].add_run("  {{ funding_source }}")
            break

    doc.save(str(OUT))
    print(f"✓ template written: {OUT}")
    print(f"  {len(BOX_TAGS)} text boxes tagged (personnel/animal filled at runtime)")


if __name__ == "__main__":
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SRC
    if not src.exists():
        sys.exit(f"Blank form not found: {src}")
    build(src)
